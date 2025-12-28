# Quiz Examination System (OOP + GUI) - Full Features
# Updated: Added Navigation Grid (1,2,3...) + Mark for Review in Student Exam
# Fixed: Listbox selection issue in TeacherFrame

import json
import os
import time
import secrets
import string
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Set, Any, Tuple

import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog

# Export libs
try:
    from docx import Document
    from openpyxl import Workbook
except ImportError:
    # Fallback or just let it crash if not installed, but better to warn
    print("Warning: python-docx or openpyxl not installed. Export features will fail.")

DATA_FILE = "quiz_data.json"
ROLES = ["Admin", "Teacher", "Student"]
ROLE_CANON = {"admin": "Admin", "teacher": "Teacher", "student": "Student"}


# -----------------------------
# Time helpers (simple, local time)
# -----------------------------
def parse_dt(s: str) -> Optional[int]:
    """Parse 'YYYY-MM-DD HH:MM' to unix timestamp (local)."""
    s = (s or "").strip()
    try:
        return int(time.mktime(time.strptime(s, "%Y-%m-%d %H:%M")))
    except Exception:
        return None


def fmt_dt(ts: int) -> str:
    try:
        return time.strftime("%Y-%m-%d %H:%M", time.localtime(int(ts)))
    except Exception:
        return "N/A"


def fmt_dt_full(ts: float) -> str:
    try:
        return time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(float(ts)))
    except Exception:
        return "N/A"


# -----------------------------
# Data Models
# -----------------------------
@dataclass
class User:
    username: str
    password: str
    role: str
    full_name: str = ""
    dob: str = ""         # YYYY-MM-DD
    student_id: str = ""  # for students


@dataclass
class Question:
    text: str
    options: List[str]              # 4 options
    correct_indices: List[int]      # can be multiple


@dataclass
class Template:
    template_id: str
    title: str
    created_by: str
    questions: List[Question]


@dataclass
class Exam:
    exam_id: str
    template_id: str
    title: str
    created_by: str
    access_code: str
    password: str
    duration_seconds: int
    allow_review: bool
    attempt_limit: int             # 0 means unlimited
    start_ts: int
    end_ts: int
    questions: List[Question]


@dataclass
class Attempt:
    attempt_id: str
    exam_id: str
    code: str
    title: str
    username: str
    full_name: str
    student_id: str
    score: float                   # max = total questions (each question max 1)
    total: int
    started_at: float
    submitted_at: float
    time_taken_seconds: int
    answers: List[List[int]]


# -----------------------------
# Scoring (partial credit)
# -----------------------------
def score_question_partial(user_sel: Set[int], correct: Set[int], points: float) -> Tuple[float, Set[int], Set[int]]:
    """
    points: points for this question
    rule:
      earned_ratio = max(0, (c - w) / k)
      k = number of correct options
      c = selected correct count
      w = selected wrong count
    """
    if not correct:
        return (0.0, set(), set())
    c = len(user_sel & correct)
    w = len(user_sel - correct)
    k = len(correct)
    earned_ratio = (c - w) / k
    earned_ratio = max(0.0, min(1.0, earned_ratio))
    earned = earned_ratio * points
    missing = correct - user_sel
    extra = user_sel - correct
    return earned, missing, extra


# -----------------------------
# Storage + migrations
# -----------------------------
class DataStore:
    def __init__(self, path: str):
        self.path = path
        self.data: Dict[str, Any] = {"users": [], "templates": [], "exams": [], "attempts": []}
        self.load()

    def load(self):
        if not os.path.exists(self.path):
            self._seed_default()
            self.save()
            return

        try:
            with open(self.path, "r", encoding="utf-8") as f:
                raw = json.load(f)
        except Exception:
            self._seed_default()
            self.save()
            return

        # ---- migrate schema from older versions ----
        # If old keys exist, keep them but make sure new keys exist.
        self.data = raw if isinstance(raw, dict) else {}
        self.data.setdefault("users", [])
        self.data.setdefault("templates", [])
        self.data.setdefault("exams", [])
        self.data.setdefault("attempts", [])

        # Normalize user roles + fields
        for u in self.data["users"]:
            u.setdefault("full_name", "")
            u.setdefault("dob", "")
            u.setdefault("student_id", "")
            # normalize role
            role = str(u.get("role", "")).strip()
            role_norm = ROLE_CANON.get(role.lower(), role)
            if role_norm not in ROLES:
                # fallback: if unknown, set Student
                role_norm = "Student"
            u["role"] = role_norm

        # migrate templates questions
        for t in self.data["templates"]:
            t.setdefault("questions", [])
            for qu in t.get("questions", []):
                if "correct_indices" not in qu and "correct_index" in qu:
                    qu["correct_indices"] = [int(qu.get("correct_index", 0))]
                qu.pop("correct_index", None)

        # migrate exams
        for e in self.data["exams"]:
            e.setdefault("attempt_limit", 0)
            e.setdefault("allow_review", False)
            e.setdefault("password", "")
            e.setdefault("duration_seconds", 0)
            e.setdefault("start_ts", 0)
            e.setdefault("end_ts", 0)
            e.setdefault("questions", [])
            for qu in e.get("questions", []):
                if "correct_indices" not in qu and "correct_index" in qu:
                    qu["correct_indices"] = [int(qu.get("correct_index", 0))]
                qu.pop("correct_index", None)

        # migrate attempts
        for a in self.data["attempts"]:
            try:
                a["score"] = float(a.get("score", 0))
            except Exception:
                a["score"] = 0.0
            a.setdefault("answers", [])

        self.save()

    def save(self):
        with open(self.path, "w", encoding="utf-8") as f:
            json.dump(self.data, f, indent=2, ensure_ascii=False)

    def _seed_default(self):
        self.data = {
            "users": [
                asdict(User(username="admin", password="admin", role="Admin")),
                asdict(User(username="teacher", password="teacher", role="Teacher",
                           full_name="Teacher One", dob="1990-01-01")),
                asdict(User(username="student", password="student", role="Student",
                           full_name="Student One", dob="2005-01-01", student_id="SV001")),
            ],
            "templates": [],
            "exams": [],
            "attempts": [],
        }

    # ---- Users ----
    def find_user(self, username: str) -> Optional[User]:
        for u in self.data["users"]:
            if u.get("username") == username:
                return User(**u)
        return None

    def list_users(self) -> List[User]:
        return [User(**u) for u in self.data["users"]]

    def add_user(self, user: User) -> bool:
        if self.find_user(user.username):
            return False
        # normalize role on add
        user.role = ROLE_CANON.get(user.role.lower(), user.role)
        if user.role not in ROLES:
            user.role = "Student"
        self.data["users"].append(asdict(user))
        self.save()
        return True

    def update_password(self, username: str, new_password: str) -> bool:
        for u in self.data["users"]:
            if u.get("username") == username:
                u["password"] = new_password
                self.save()
                return True
        return False

    def update_profile(self, username: str, full_name: str, dob: str, student_id: str) -> bool:
        for u in self.data["users"]:
            if u.get("username") == username:
                u["full_name"] = full_name
                u["dob"] = dob
                u["student_id"] = student_id
                self.save()
                return True
        return False

    # ---- IDs/Codes ----
    def new_template_id(self) -> str:
        return f"TP{int(time.time() * 1000)}"

    def new_exam_id(self) -> str:
        return f"EX{int(time.time() * 1000)}"

    def new_attempt_id(self) -> str:
        return f"AT{int(time.time() * 1000)}{secrets.randbelow(1000)}"

    def _all_codes(self) -> Set[str]:
        return {(e.get("access_code") or "").upper() for e in self.data["exams"] if e.get("access_code")}

    def new_unique_code(self, length: int = 8) -> str:
        alphabet = string.ascii_uppercase + string.digits
        existing = self._all_codes()
        while True:
            code = "".join(secrets.choice(alphabet) for _ in range(length))
            if code not in existing:
                return code

    # ---- Templates ----
    def add_template(self, t: Template):
        self.data["templates"].append(self._template_to_dict(t))
        self.save()

    def list_templates(self) -> List[Template]:
        return [self._dict_to_template(x) for x in self.data["templates"]]

    def list_templates_by_teacher(self, teacher: str) -> List[Template]:
        return [t for t in self.list_templates() if t.created_by == teacher]

    def get_template(self, template_id: str) -> Optional[Template]:
        for t in self.data["templates"]:
            if t.get("template_id") == template_id:
                return self._dict_to_template(t)
        return None

    def delete_template(self, template_id: str) -> bool:
        before = len(self.data["templates"])
        self.data["templates"] = [t for t in self.data["templates"] if t.get("template_id") != template_id]
        after = len(self.data["templates"])
        if after == before:
            return False
        self.save()
        return True

    def has_exam_from_template(self, template_id: str) -> bool:
        return any(e.get("template_id") == template_id for e in self.data["exams"])

    def _template_to_dict(self, t: Template) -> Dict[str, Any]:
        return {
            "template_id": t.template_id,
            "title": t.title,
            "created_by": t.created_by,
            "questions": [asdict(q) for q in t.questions],
        }

    def _dict_to_template(self, d: Dict[str, Any]) -> Template:
        qs = []
        for q in d.get("questions", []):
            qs.append(Question(text=q["text"], options=q["options"], correct_indices=q.get("correct_indices", [])))
        return Template(
            template_id=d.get("template_id", ""),
            title=d.get("title", ""),
            created_by=d.get("created_by", ""),
            questions=qs
        )

    # ---- Exams ----
    def add_exam(self, e: Exam):
        self.data["exams"].append(self._exam_to_dict(e))
        self.save()

    def list_exams(self) -> List[Exam]:
        return [self._dict_to_exam(x) for x in self.data["exams"]]

    def list_exams_by_teacher(self, teacher: str) -> List[Exam]:
        return [e for e in self.list_exams() if e.created_by == teacher]

    def get_exam_by_code(self, code: str) -> Optional[Exam]:
        code = (code or "").strip().upper()
        for e in self.data["exams"]:
            if (e.get("access_code") or "").upper() == code:
                return self._dict_to_exam(e)
        return None

    def get_exam(self, exam_id: str) -> Optional[Exam]:
        for e in self.data["exams"]:
            if e.get("exam_id") == exam_id:
                return self._dict_to_exam(e)
        return None

    def delete_exam(self, exam_id: str) -> bool:
        before = len(self.data["exams"])
        self.data["exams"] = [e for e in self.data["exams"] if e.get("exam_id") != exam_id]
        after = len(self.data["exams"])
        if after == before:
            return False
        self.save()
        return True

    def _exam_to_dict(self, e: Exam) -> Dict[str, Any]:
        return {
            "exam_id": e.exam_id,
            "template_id": e.template_id,
            "title": e.title,
            "created_by": e.created_by,
            "access_code": e.access_code,
            "password": e.password,
            "duration_seconds": e.duration_seconds,
            "allow_review": e.allow_review,
            "attempt_limit": e.attempt_limit,
            "start_ts": e.start_ts,
            "end_ts": e.end_ts,
            "questions": [asdict(q) for q in e.questions],
        }

    def _dict_to_exam(self, d: Dict[str, Any]) -> Exam:
        qs = []
        for q in d.get("questions", []):
            qs.append(Question(text=q["text"], options=q["options"], correct_indices=q.get("correct_indices", [])))
        return Exam(
            exam_id=d.get("exam_id", ""),
            template_id=d.get("template_id", ""),
            title=d.get("title", ""),
            created_by=d.get("created_by", ""),
            access_code=d.get("access_code", ""),
            password=d.get("password", "") or "",
            duration_seconds=int(d.get("duration_seconds", 0) or 0),
            allow_review=bool(d.get("allow_review", False)),
            attempt_limit=int(d.get("attempt_limit", 0) or 0),
            start_ts=int(d.get("start_ts", 0) or 0),
            end_ts=int(d.get("end_ts", 0) or 0),
            questions=qs
        )

    # ---- Attempts ----
    def add_attempt(self, a: Attempt):
        self.data["attempts"].append(asdict(a))
        self.save()

    def list_attempts_for_user(self, username: str) -> List[Attempt]:
        out = [Attempt(**x) for x in self.data["attempts"] if x.get("username") == username]
        out.sort(key=lambda z: z.submitted_at, reverse=True)
        return out

    def list_attempts_for_exam(self, exam_id: str) -> List[Attempt]:
        out = [Attempt(**x) for x in self.data["attempts"] if x.get("exam_id") == exam_id]
        out.sort(key=lambda z: z.submitted_at, reverse=True)
        return out

    def count_attempts_for_user_exam(self, username: str, exam_id: str) -> int:
        return sum(1 for x in self.data["attempts"] if x.get("username") == username and x.get("exam_id") == exam_id)

    def delete_attempts_for_exam(self, exam_id: str) -> int:
        before = len(self.data["attempts"])
        self.data["attempts"] = [a for a in self.data["attempts"] if a.get("exam_id") != exam_id]
        after = len(self.data["attempts"])
        deleted = before - after
        if deleted:
            self.save()
        return deleted


# -----------------------------
# Export helpers
# -----------------------------
def export_template_to_word(t: Template, filepath: str, include_answers: bool = True):
    try:
        doc = Document()
        doc.add_heading(f"Template: {t.title}", level=1)
        doc.add_paragraph(f"Template ID: {t.template_id}")
        doc.add_paragraph(f"Created by: {t.created_by}")
        doc.add_paragraph(f"Questions: {len(t.questions)}")

        for i, q in enumerate(t.questions, start=1):
            doc.add_heading(f"Q{i}: {q.text}", level=2)
            for oi, opt in enumerate(q.options, start=1):
                doc.add_paragraph(f"{oi}. {opt}", style="List Bullet")
            if include_answers:
                correct = ", ".join(str(x + 1) for x in sorted(q.correct_indices))
                doc.add_paragraph(f"Correct: {correct}")
        doc.save(filepath)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export docx: {e}")


def export_exam_to_word(e: Exam, filepath: str, include_answers: bool = True):
    try:
        doc = Document()
        doc.add_heading(f"Exam: {e.title}", level=1)
        doc.add_paragraph(f"Exam ID: {e.exam_id}")
        doc.add_paragraph(f"Code: {e.access_code}")
        doc.add_paragraph(f"Window: {fmt_dt(e.start_ts)} -> {fmt_dt(e.end_ts)}")
        doc.add_paragraph(f"Duration: {max(1, e.duration_seconds // 60)} minutes")
        doc.add_paragraph(f"Password: {'set' if e.password else 'none'}")
        doc.add_paragraph(f"Allow review: {'yes' if e.allow_review else 'no'}")
        doc.add_paragraph(f"Attempt limit: {e.attempt_limit} (0=unlimited)")
        doc.add_paragraph(f"From template: {e.template_id}")

        for i, q in enumerate(e.questions, start=1):
            doc.add_heading(f"Q{i}: {q.text}", level=2)
            for oi, opt in enumerate(q.options, start=1):
                doc.add_paragraph(f"{oi}. {opt}", style="List Bullet")
            if include_answers:
                correct = ", ".join(str(x + 1) for x in sorted(q.correct_indices))
                doc.add_paragraph(f"Correct: {correct}")
        doc.save(filepath)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export docx: {e}")


def export_exam_results_to_excel(e: Exam, attempts: List[Attempt], filepath: str):
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Results"

        ws.append([
            "Exam Title", "Exam ID", "Code",
            "Username", "Full Name", "Student ID",
            "Score (/10)", "Raw Score", "Total Questions",
            "Time Taken (sec)", "Submitted At"
        ])

        total_q = max(1, len(e.questions))
        for a in attempts:
            score10 = (a.score / max(1, a.total)) * 10.0
            ws.append([
                e.title, e.exam_id, e.access_code,
                a.username, a.full_name, a.student_id,
                round(score10, 2), round(a.score, 4), a.total,
                a.time_taken_seconds, fmt_dt_full(a.submitted_at)
            ])
        wb.save(filepath)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export xlsx: {e}")


def export_attempt_to_word(e: Exam, a: Attempt, filepath: str):
    try:
        doc = Document()
        doc.add_heading("Attempt Report", level=1)
        doc.add_paragraph(f"Exam: {e.title}")
        doc.add_paragraph(f"Exam ID: {e.exam_id} | Code: {e.access_code}")
        doc.add_paragraph(f"Student: {a.full_name} | Username: {a.username} | Student ID: {a.student_id}")
        doc.add_paragraph(f"Started: {fmt_dt_full(a.started_at)}")
        doc.add_paragraph(f"Submitted: {fmt_dt_full(a.submitted_at)}")
        doc.add_paragraph(f"Time taken: {a.time_taken_seconds} sec")

        total_q = max(1, len(e.questions))
        score10 = (a.score / total_q) * 10.0
        doc.add_paragraph(f"Score: {score10:.2f}/10 (raw {a.score:.4f}/{total_q})")

        answers = a.answers or []
        points_per_q = 10.0 / total_q

        for i, q in enumerate(e.questions):
            user_sel = set(answers[i]) if i < len(answers) else set()
            correct = set(q.correct_indices)
            earned, missing, extra = score_question_partial(user_sel, correct, points=points_per_q)

            doc.add_heading(f"Q{i+1}: {q.text}", level=2)
            doc.add_paragraph(f"Earned: {earned:.2f}/{points_per_q:.2f}")
            for oi, opt in enumerate(q.options, start=1):
                mu = "[x]" if (oi - 1) in user_sel else "[ ]"
                mc = "(correct)" if (oi - 1) in correct else ""
                doc.add_paragraph(f"{mu} {oi}. {opt} {mc}")
            doc.add_paragraph("Missing correct: " + (", ".join(str(x+1) for x in sorted(missing)) if missing else "none"))
            doc.add_paragraph("Extra wrong: " + (", ".join(str(x+1) for x in sorted(extra)) if extra else "none"))

        doc.save(filepath)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export docx: {e}")


# -----------------------------
# UI helpers
# -----------------------------
def info(msg: str):
    messagebox.showinfo("Info", msg)


def err(msg: str):
    messagebox.showerror("Error", msg)


class Header(ttk.Frame):
    def __init__(self, parent, title: str, subtitle: str = ""):
        super().__init__(parent)
        ttk.Label(self, text=title, font=("Segoe UI", 18, "bold")).pack(anchor="w")
        if subtitle:
            ttk.Label(self, text=subtitle).pack(anchor="w", pady=(2, 0))


# -----------------------------
# App Core (reload safe)
# -----------------------------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Quiz Examination System")
        self.geometry("1180x760")
        self.minsize(1180, 760)

        self.store = DataStore(DATA_FILE)
        self.current_user: Optional[User] = None
        self.current_frame_name: str = "LoginFrame"

        self.container = ttk.Frame(self, padding=12)
        self.container.pack(fill="both", expand=True)
        self.container.rowconfigure(0, weight=1)
        self.container.columnconfigure(0, weight=1)

        self.frames: Dict[str, ttk.Frame] = {}
        self._init_frames()
        self.show_frame("LoginFrame")

    def _init_frames(self):
        for F in (
            LoginFrame, AdminFrame, TeacherFrame, StudentFrame,
            ExamTakeFrame, ReviewFrame, TeacherAttemptFrame,
            TemplatePreviewFrame, ExamPreviewFrame
        ):
            frame = F(self.container, self)
            self.frames[F.__name__] = frame
            frame.grid(row=0, column=0, sticky="nsew")

    def show_frame(self, name: str):
        self.current_frame_name = name
        frame = self.frames[name]
        if hasattr(frame, "on_show"):
            frame.on_show()
        frame.tkraise()

    def reload_data(self):
        self.store.load()
        if self.current_user:
            u = self.store.find_user(self.current_user.username)
            if u:
                self.current_user = u
        cur = self.frames.get(self.current_frame_name)
        if cur and hasattr(cur, "on_show"):
            cur.on_show()

    def logout(self):
        tf: ExamTakeFrame = self.frames["ExamTakeFrame"]  # type: ignore
        tf.stop_timer()
        self.current_user = None
        self.show_frame("LoginFrame")


# -----------------------------
# Login Frame
# -----------------------------
class LoginFrame(ttk.Frame):
    def __init__(self, parent, app: App):
        super().__init__(parent)
        self.app = app

        Header(
            self,
            "Quiz Examination System",
            "Login. Default: admin/admin, teacher/teacher, student/student."
        ).pack(fill="x", pady=(0, 16))

        form = ttk.Frame(self)
        form.pack(pady=26)

        ttk.Label(form, text="Role:").grid(row=0, column=0, sticky="e", padx=8, pady=8)
        self.role_var = tk.StringVar(value=ROLES[0])
        ttk.Combobox(form, textvariable=self.role_var, values=ROLES, state="readonly", width=22)\
            .grid(row=0, column=1, sticky="w", padx=8, pady=8)

        ttk.Label(form, text="Username:").grid(row=1, column=0, sticky="e", padx=8, pady=8)
        self.user_var = tk.StringVar()
        self.user_entry = ttk.Entry(form, textvariable=self.user_var, width=26)
        self.user_entry.grid(row=1, column=1, sticky="w", padx=8, pady=8)

        ttk.Label(form, text="Password:").grid(row=2, column=0, sticky="e", padx=8, pady=8)
        self.pass_var = tk.StringVar()
        self.pass_entry = ttk.Entry(form, textvariable=self.pass_var, show="*", width=26)
        self.pass_entry.grid(row=2, column=1, sticky="w", padx=8, pady=8)

        btns = ttk.Frame(self)
        btns.pack(pady=8)
        ttk.Button(btns, text="Login", command=self.do_login).pack(side="left", padx=6)
        ttk.Button(btns, text="Clear", command=self.clear).pack(side="left", padx=6)

        self.user_entry.focus_set()
        self.bind_all("<Return>", self._enter_login)

    def _enter_login(self, event):
        if self.winfo_ismapped():
            self.do_login()

    def clear(self):
        self.user_var.set("")
        self.pass_var.set("")
        self.user_entry.focus_set()

    def do_login(self):
        role = self.role_var.get().strip()
        username = self.user_var.get().strip()
        password = self.pass_var.get()

        if role not in ROLES or not username or not password:
            err("Role, username, and password are needed.")
            return

        u = self.app.store.find_user(username)
        if not u:
            err("User not found.")
            return
        if u.password != password:
            err("Password is not correct.")
            return

        # Ensure role compare uses canonical set
        u.role = ROLE_CANON.get(u.role.lower(), u.role)
        if u.role != role:
            err(f"Role mismatch.\nAccount role: {u.role}\nYou selected: {role}")
            return

        self.app.current_user = u
        info(f"Login ok. Welcome {u.username} ({u.role}).")
        if u.role == "Admin":
            self.app.show_frame("AdminFrame")
        elif u.role == "Teacher":
            self.app.show_frame("TeacherFrame")
        else:
            self.app.show_frame("StudentFrame")


# -----------------------------
# Admin Frame
# -----------------------------
class AdminFrame(ttk.Frame):
    def __init__(self, parent, app: App):
        super().__init__(parent)
        self.app = app
        Header(self, "Admin Panel", "Admin manages users.").pack(fill="x", pady=(0, 10))

        top = ttk.Frame(self)
        top.pack(fill="x")
        ttk.Button(top, text="Reload", command=self.app.reload_data).pack(side="right", padx=6)
        ttk.Button(top, text="Logout", command=self.app.logout).pack(side="right")

        main = ttk.Frame(self)
        main.pack(fill="both", expand=True, pady=10)

        left = ttk.LabelFrame(main, text="Create User")
        left.pack(side="left", fill="both", expand=True, padx=(0, 8))
        right = ttk.LabelFrame(main, text="Users / Reset password")
        right.pack(side="left", fill="both", expand=True, padx=(8, 0))

        self.new_user = tk.StringVar()
        self.new_pass = tk.StringVar()
        self.new_role = tk.StringVar(value="Student")

        f = ttk.Frame(left, padding=10)
        f.pack(fill="x")
        ttk.Label(f, text="Username:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        ttk.Entry(f, textvariable=self.new_user, width=28).grid(row=0, column=1, padx=6, pady=6, sticky="w")
        ttk.Label(f, text="Password:").grid(row=1, column=0, sticky="e", padx=6, pady=6)
        ttk.Entry(f, textvariable=self.new_pass, width=28).grid(row=1, column=1, padx=6, pady=6, sticky="w")
        ttk.Label(f, text="Role:").grid(row=2, column=0, sticky="e", padx=6, pady=6)
        ttk.Combobox(f, textvariable=self.new_role, values=ROLES, state="readonly", width=26)\
            .grid(row=2, column=1, padx=6, pady=6, sticky="w")
        ttk.Button(f, text="Create", command=self.create_user).grid(row=3, column=0, columnspan=2, pady=(10, 4))

        self.reset_user = tk.StringVar()
        self.reset_pass = tk.StringVar()
        rr = ttk.Frame(right, padding=10)
        rr.pack(fill="x")
        ttk.Label(rr, text="Username:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        ttk.Entry(rr, textvariable=self.reset_user, width=26).grid(row=0, column=1, sticky="w", padx=6, pady=6)
        ttk.Label(rr, text="New password:").grid(row=1, column=0, sticky="e", padx=6, pady=6)
        ttk.Entry(rr, textvariable=self.reset_pass, width=26).grid(row=1, column=1, sticky="w", padx=6, pady=6)
        ttk.Button(rr, text="Reset", command=self.reset_password).grid(row=2, column=0, columnspan=2, pady=(10, 4))

        self.users_box = tk.Listbox(right, height=18)
        self.users_box.pack(fill="both", expand=True, padx=10, pady=(6, 10))

    def on_show(self):
        if not self.app.current_user or self.app.current_user.role != "Admin":
            err("You need Admin role.")
            self.app.logout()
            return
        self.refresh_users()

    def refresh_users(self):
        self.users_box.delete(0, tk.END)
        for u in self.app.store.list_users():
            self.users_box.insert(tk.END, f"{u.username} | {u.role} | {u.full_name} | {u.dob} | {u.student_id}")

    def create_user(self):
        username = self.new_user.get().strip()
        password = self.new_pass.get()
        role = self.new_role.get().strip()
        if not username or not password or role not in ROLES:
            err("Username, password, and role are needed.")
            return
        ok = self.app.store.add_user(User(username=username, password=password, role=role))
        if not ok:
            err("Username exists.")
            return
        info("User created.")
        self.new_user.set("")
        self.new_pass.set("")
        self.new_role.set("Student")
        self.refresh_users()

    def reset_password(self):
        username = self.reset_user.get().strip()
        new_password = self.reset_pass.get()
        if not username or not new_password:
            err("Username and new password are needed.")
            return
        ok = self.app.store.update_password(username, new_password)
        if not ok:
            err("User not found.")
            return
        info("Password updated.")
        self.reset_user.set("")
        self.reset_pass.set("")
        self.refresh_users()


# -----------------------------
# Teacher Frame (delete + export)
# -----------------------------
class TeacherFrame(ttk.Frame):
    def __init__(self, parent, app: App):
        super().__init__(parent)
        self.app = app
        self.selected_exam_id: Optional[str] = None

        Header(
            self,
            "Teacher Panel",
            "Templates = warehouse. Exams = published for students (code + time window)."
        ).pack(fill="x", pady=(0, 10))

        top = ttk.Frame(self)
        top.pack(fill="x")
        ttk.Button(top, text="Reload", command=self.app.reload_data).pack(side="right", padx=6)
        ttk.Button(top, text="Logout", command=self.app.logout).pack(side="right")

        main = ttk.Frame(self)
        main.pack(fill="both", expand=True, pady=10)

        left = ttk.LabelFrame(main, text="Template builder (warehouse)")
        left.pack(side="left", fill="both", expand=True, padx=(0, 8))

        right = ttk.LabelFrame(main, text="Publish / Exams / Attempts")
        right.pack(side="left", fill="both", expand=True, padx=(8, 0))

        # builder
        self.tpl_title = tk.StringVar()
        self.q_text = tk.StringVar()
        self.opt_vars = [tk.StringVar() for _ in range(4)]
        self.correct_vars = [tk.BooleanVar(value=False) for _ in range(4)]
        self._temp_questions: List[Question] = []

        bf = ttk.Frame(left, padding=10)
        bf.pack(fill="x")
        ttk.Label(bf, text="Template title:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        ttk.Entry(bf, textvariable=self.tpl_title, width=34).grid(row=0, column=1, sticky="w", padx=6, pady=6)

        ttk.Separator(left).pack(fill="x", padx=10, pady=6)

        qf = ttk.Frame(left, padding=10)
        qf.pack(fill="x")

        ttk.Label(qf, text="Question text:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        ttk.Entry(qf, textvariable=self.q_text, width=34).grid(row=0, column=1, sticky="w", padx=6, pady=6)

        for i in range(4):
            r = 1 + i
            ttk.Label(qf, text=f"Option {i+1}:").grid(row=r, column=0, sticky="e", padx=6, pady=4)
            row = ttk.Frame(qf)
            row.grid(row=r, column=1, sticky="w", padx=6, pady=4)
            ttk.Entry(row, textvariable=self.opt_vars[i], width=28).pack(side="left")
            ttk.Checkbutton(row, text="Correct", variable=self.correct_vars[i]).pack(side="left", padx=8)

        br = ttk.Frame(left, padding=10)
        br.pack(fill="x")
        ttk.Button(br, text="Add question", command=self.add_question).pack(side="left", padx=6)
        ttk.Button(br, text="Save template", command=self.save_template).pack(side="left", padx=6)
        ttk.Button(br, text="Clear", command=self.clear_builder).pack(side="left", padx=6)

        self.temp_list = tk.Listbox(left, height=10)
        self.temp_list.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        # right panel
        rf = ttk.Frame(right, padding=10)
        rf.pack(fill="both", expand=True)

        ttk.Label(rf, text="My templates (select):").pack(anchor="w")
        self.tpl_list = tk.Listbox(rf, height=7)
        self.tpl_list.pack(fill="x", pady=(6, 8))

        tpl_btn = ttk.Frame(rf)
        tpl_btn.pack(fill="x", pady=(0, 8))
        ttk.Button(tpl_btn, text="Preview template", command=self.preview_selected_template).pack(side="left")
        ttk.Button(tpl_btn, text="Delete template", command=self.delete_selected_template).pack(side="left", padx=6)
        ttk.Button(tpl_btn, text="Export template (Word)", command=self.export_selected_template_word).pack(side="right")

        pub = ttk.LabelFrame(rf, text="Publish exam from selected template")
        pub.pack(fill="x", pady=(0, 8))

        self.pub_start = tk.StringVar(value=time.strftime("%Y-%m-%d 08:00", time.localtime()))
        self.pub_end = tk.StringVar(value=time.strftime("%Y-%m-%d 23:00", time.localtime()))
        self.pub_duration = tk.IntVar(value=30)
        self.pub_use_pass = tk.BooleanVar(value=False)
        self.pub_pass = tk.StringVar()
        self.pub_allow_review = tk.BooleanVar(value=False)
        self.pub_attempt_limit = tk.IntVar(value=1)

        pf = ttk.Frame(pub, padding=8)
        pf.pack(fill="x")

        ttk.Label(pf, text="Start (YYYY-MM-DD HH:MM):").grid(row=0, column=0, sticky="e", padx=6, pady=4)
        ttk.Entry(pf, textvariable=self.pub_start, width=20).grid(row=0, column=1, sticky="w", padx=6, pady=4)
        ttk.Label(pf, text="End (YYYY-MM-DD HH:MM):").grid(row=0, column=2, sticky="e", padx=6, pady=4)
        ttk.Entry(pf, textvariable=self.pub_end, width=20).grid(row=0, column=3, sticky="w", padx=6, pady=4)

        ttk.Label(pf, text="Duration (minutes):").grid(row=1, column=0, sticky="e", padx=6, pady=4)
        ttk.Spinbox(pf, from_=1, to=240, textvariable=self.pub_duration, width=6)\
            .grid(row=1, column=1, sticky="w", padx=6, pady=4)

        ttk.Label(pf, text="Attempt limit (0=unlimited):").grid(row=1, column=2, sticky="e", padx=6, pady=4)
        ttk.Spinbox(pf, from_=0, to=50, textvariable=self.pub_attempt_limit, width=6)\
            .grid(row=1, column=3, sticky="w", padx=6, pady=4)

        pr = ttk.Frame(pf)
        pr.grid(row=2, column=0, columnspan=4, sticky="w", pady=(2, 2))
        ttk.Checkbutton(pr, text="Use password (may)", variable=self.pub_use_pass, command=self._toggle_pub_pass).pack(side="left")
        ttk.Label(pr, text="Password:").pack(side="left", padx=(10, 0))
        self.pub_pass_entry = ttk.Entry(pr, textvariable=self.pub_pass, width=16, show="*")
        self.pub_pass_entry.pack(side="left", padx=6)
        ttk.Checkbutton(pr, text="Allow review (may)", variable=self.pub_allow_review).pack(side="left", padx=(12, 0))

        btnpub = ttk.Frame(pub, padding=(8, 0, 8, 8))
        btnpub.pack(fill="x")
        ttk.Button(btnpub, text="Publish", command=self.publish_exam).pack(side="left")

        ttk.Separator(rf).pack(fill="x", pady=6)

        ttk.Label(rf, text="My exams (select):").pack(anchor="w")
        # FIXED: exportselection=False to prevent losing selection when clicking other lists
        self.exam_list = tk.Listbox(rf, height=7, exportselection=False)
        self.exam_list.pack(fill="x", pady=(6, 8))
        self.exam_list.bind("<<ListboxSelect>>", self._on_exam_select)

        exam_btn = ttk.Frame(rf)
        exam_btn.pack(fill="x", pady=(0, 8))
        ttk.Button(exam_btn, text="Preview exam", command=self.preview_selected_exam).pack(side="left")
        ttk.Button(exam_btn, text="Delete exam", command=self.delete_selected_exam).pack(side="left", padx=6)
        ttk.Button(exam_btn, text="Delete attempts", command=self.delete_attempts_for_selected_exam).pack(side="left", padx=6)
        ttk.Button(exam_btn, text="Export exam (Word)", command=self.export_selected_exam_word).pack(side="right")
        ttk.Button(exam_btn, text="Export results (Excel)", command=self.export_selected_exam_results_excel).pack(side="right", padx=6)

        ttk.Label(rf, text="Attempts (double click to open):").pack(anchor="w", pady=(8, 0))
        # FIXED: exportselection=False
        self.attempt_list = tk.Listbox(rf, height=9, exportselection=False)
        self.attempt_list.pack(fill="both", expand=True, pady=(6, 0))
        self.attempt_list.bind("<Double-Button-1>", lambda e: self.view_attempt_details())
        self.selected_attempt_id: str = ""

        self.attempt_list.bind("<<ListboxSelect>>", self._on_attempt_select)
        self.attempt_list.bind("<Button-1>", self._on_attempt_click)
        self.attempt_list.bind("<Double-Button-1>", lambda e: self.view_attempt_details())

        att_btn = ttk.Frame(rf)
        att_btn.pack(fill="x", pady=(6, 0))
        ttk.Button(att_btn, text="View attempt details", command=self.view_attempt_details).pack(side="left")
        ttk.Button(att_btn, text="Export attempt (Word)", command=self.export_selected_attempt_word).pack(side="left", padx=6)

        self._attempt_id_by_index: Dict[int, str] = {}
        self._toggle_pub_pass()

    def _toggle_pub_pass(self):
        if self.pub_use_pass.get():
            self.pub_pass_entry.config(state="normal")
        else:
            self.pub_pass_entry.config(state="disabled")
            self.pub_pass.set("")

    def on_show(self):
        if not self.app.current_user or self.app.current_user.role != "Teacher":
            err("You need Teacher role.")
            self.app.logout()
            return
        self.refresh_templates()
        self.refresh_exams()
        self.attempt_list.delete(0, tk.END)
        self._attempt_id_by_index.clear()

    def refresh_templates(self):
        self.tpl_list.delete(0, tk.END)
        teacher = self.app.current_user.username
        for t in self.app.store.list_templates_by_teacher(teacher):
            self.tpl_list.insert(tk.END, f"{t.template_id} | {t.title} | questions:{len(t.questions)}")

    def refresh_exams(self):
        self.exam_list.delete(0, tk.END)
        teacher = self.app.current_user.username
        now = int(time.time())
        for e in self.app.store.list_exams_by_teacher(teacher):
            status = "OPEN" if (e.start_ts <= now <= e.end_ts) else ("WAIT" if now < e.start_ts else "CLOSED")
            self.exam_list.insert(
                tk.END,
                f"{e.exam_id} | {e.title} | code:{e.access_code} | {status} | {fmt_dt(e.start_ts)} -> {fmt_dt(e.end_ts)} | limit:{e.attempt_limit}"
            )

    # ---- builder ----
    def add_question(self):
        text = self.q_text.get().strip()
        options = [v.get().strip() for v in self.opt_vars]
        correct_indices = [i for i, b in enumerate(self.correct_vars) if b.get()]
        if not text:
            err("Question text is needed.")
            return
        if any(not o for o in options):
            err("All 4 options are needed.")
            return
        if len(correct_indices) == 0:
            err("Need at least 1 correct option.")
            return
        q = Question(text=text, options=options, correct_indices=correct_indices)
        self._temp_questions.append(q)
        self.temp_list.insert(tk.END, f"{len(self._temp_questions)}. {text} (correct: {','.join(str(x+1) for x in correct_indices)})")
        self.q_text.set("")
        for v in self.opt_vars:
            v.set("")
        for b in self.correct_vars:
            b.set(False)

    def clear_builder(self):
        self.tpl_title.set("")
        self.q_text.set("")
        for v in self.opt_vars:
            v.set("")
        for b in self.correct_vars:
            b.set(False)
        self._temp_questions.clear()
        self.temp_list.delete(0, tk.END)

    def save_template(self):
        title = self.tpl_title.get().strip()
        if not title:
            err("Template title is needed.")
            return
        if len(self._temp_questions) == 0:
            err("Need at least 1 question.")
            return
        t = Template(
            template_id=self.app.store.new_template_id(),
            title=title,
            created_by=self.app.current_user.username,
            questions=list(self._temp_questions)
        )
        self.app.store.add_template(t)
        info("Template saved into warehouse.")
        self.clear_builder()
        self.refresh_templates()
        pv: TemplatePreviewFrame = self.app.frames["TemplatePreviewFrame"]  # type: ignore
        pv.load_template(t.template_id, back_to="TeacherFrame")
        self.app.show_frame("TemplatePreviewFrame")

    # ---- selection helpers ----
    def _selected_template_id(self) -> Optional[str]:
        sel = self.tpl_list.curselection()
        if not sel:
            return None
        return self.tpl_list.get(sel[0]).split("|")[0].strip()

    def _selected_exam_id(self) -> Optional[str]:
        sel = self.exam_list.curselection()
        if not sel:
            return None
        return self.exam_list.get(sel[0]).split("|")[0].strip()

    # ---- delete template ----
    def delete_selected_template(self):
        tid = self._selected_template_id()
        if not tid:
            err("Select a template first.")
            return
        tpl = self.app.store.get_template(tid)
        if not tpl:
            err("Template not found.")
            return
        # if there are published exams from this template
        if self.app.store.has_exam_from_template(tid):
            if not messagebox.askyesno("Confirm", "This template has published exams.\nDelete template anyway? (Exams will remain)"):
                return
        if not messagebox.askyesno("Confirm", f"Delete template '{tpl.title}'?"):
            return
        ok = self.app.store.delete_template(tid)
        if ok:
            info("Template deleted.")
            self.refresh_templates()
        else:
            err("Delete failed.")

    # ---- template preview/export ----
    def preview_selected_template(self):
        tid = self._selected_template_id()
        if not tid:
            err("Select a template first.")
            return
        pv: TemplatePreviewFrame = self.app.frames["TemplatePreviewFrame"]  # type: ignore
        pv.load_template(tid, back_to="TeacherFrame")
        self.app.show_frame("TemplatePreviewFrame")

    def export_selected_template_word(self):
        tid = self._selected_template_id()
        if not tid:
            err("Select a template first.")
            return
        tpl = self.app.store.get_template(tid)
        if not tpl:
            err("Template not found.")
            return
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"template_{tpl.template_id}.docx"
        )
        if not filepath:
            return
        include_ans = messagebox.askyesno("Include answers?", "Do you want to include correct answers in Word?")
        export_template_to_word(tpl, filepath, include_answers=include_ans)
        info("Exported template to Word.")

    # ---- publish ----
    def publish_exam(self):
        tid = self._selected_template_id()
        if not tid:
            err("Select a template to publish.")
            return
        tpl = self.app.store.get_template(tid)
        if not tpl:
            err("Template not found. Please reload.")
            return

        start_ts = parse_dt(self.pub_start.get())
        end_ts = parse_dt(self.pub_end.get())
        if start_ts is None or end_ts is None:
            err("Start/End format: YYYY-MM-DD HH:MM")
            return
        if end_ts <= start_ts:
            err("End must be after Start.")
            return

        dur_min = int(self.pub_duration.get())
        if dur_min <= 0:
            err("Duration minutes must be >= 1.")
            return

        attempt_limit = int(self.pub_attempt_limit.get())
        if attempt_limit < 0:
            attempt_limit = 0

        password = ""
        if self.pub_use_pass.get():
            password = self.pub_pass.get()
            if not password:
                err("Password is needed if you use password.")
                return

        code = self.app.store.new_unique_code(8)
        e = Exam(
            exam_id=self.app.store.new_exam_id(),
            template_id=tpl.template_id,
            title=tpl.title,
            created_by=self.app.current_user.username,
            access_code=code,
            password=password,
            duration_seconds=dur_min * 60,
            allow_review=bool(self.pub_allow_review.get()),
            attempt_limit=attempt_limit,
            start_ts=int(start_ts),
            end_ts=int(end_ts),
            questions=list(tpl.questions)
        )
        self.app.store.add_exam(e)
        info(f"Exam published.\nCODE: {code}")
        self.refresh_exams()
        pv: ExamPreviewFrame = self.app.frames["ExamPreviewFrame"]  # type: ignore
        pv.load_exam(e.exam_id, back_to="TeacherFrame")
        self.app.show_frame("ExamPreviewFrame")

    # ---- exam preview/delete/export ----
    def preview_selected_exam(self):
        eid = self._selected_exam_id()
        if not eid:
            err("Select an exam first.")
            return
        pv: ExamPreviewFrame = self.app.frames["ExamPreviewFrame"]  # type: ignore
        pv.load_exam(eid, back_to="TeacherFrame")
        self.app.show_frame("ExamPreviewFrame")

    def delete_selected_exam(self):
        eid = self._selected_exam_id()
        if not eid:
            err("Select an exam first.")
            return
        e = self.app.store.get_exam(eid)
        if not e:
            err("Exam not found.")
            return
        if not messagebox.askyesno("Confirm", f"Delete exam '{e.title}' (code {e.access_code})?\nAttempts will NOT be deleted automatically."):
            return
        ok = self.app.store.delete_exam(eid)
        if ok:
            info("Exam deleted.")
            self.refresh_exams()
            self.attempt_list.delete(0, tk.END)
            self._attempt_id_by_index.clear()
        else:
            err("Delete failed.")

    def delete_attempts_for_selected_exam(self):
        eid = self._selected_exam_id()
        if not eid:
            err("Select an exam first.")
            return
        if not messagebox.askyesno("Confirm", "Delete ALL attempts for this exam?"):
            return
        deleted = self.app.store.delete_attempts_for_exam(eid)
        info(f"Deleted attempts: {deleted}")
        self._on_exam_select()

    def export_selected_exam_word(self):
        eid = self._selected_exam_id()
        if not eid:
            err("Select an exam first.")
            return
        e = self.app.store.get_exam(eid)
        if not e:
            err("Exam not found.")
            return
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"exam_{e.exam_id}_{e.access_code}.docx"
        )
        if not filepath:
            return
        include_ans = messagebox.askyesno("Include answers?", "Do you want to include correct answers in Word?")
        export_exam_to_word(e, filepath, include_answers=include_ans)
        info("Exported exam to Word.")

    def export_selected_exam_results_excel(self):
        eid = self._selected_exam_id()
        if not eid:
            err("Select an exam first.")
            return
        e = self.app.store.get_exam(eid)
        if not e:
            err("Exam not found.")
            return
        attempts = self.app.store.list_attempts_for_exam(eid)
        filepath = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel Workbook", "*.xlsx")],
            initialfile=f"results_{e.exam_id}_{e.access_code}.xlsx"
        )
        if not filepath:
            return
        export_exam_results_to_excel(e, attempts, filepath)
        info("Exported results to Excel.")

    # ---- attempts list ----
    def _on_exam_select(self, event=None):
        eid = self._selected_exam_id()
        self.selected_exam_id = eid

        self.attempt_list.delete(0, tk.END)
        self._attempt_id_by_index.clear()

        if not eid:
            return

        attempts = self.app.store.list_attempts_for_exam(eid)
        if not attempts:
            self.attempt_list.insert(tk.END, "(No attempts yet)")
            return

        for idx, a in enumerate(attempts):
            took = f"{a.time_taken_seconds // 60:02d}:{a.time_taken_seconds % 60:02d}"
            submit_t = fmt_dt_full(a.submitted_at)
            score10 = (a.score / max(1, a.total)) * 10.0
            name = a.full_name or "(no name)"
            self.attempt_list.insert(
                tk.END,
                f"{name} | {a.username} | {score10:.2f}/10 | took {took} | {submit_t}"
            )
            self._attempt_id_by_index[idx] = a.attempt_id
        
        # Auto select first
        if attempts:
             self.attempt_list.selection_set(0)
             self.attempt_list.activate(0)
             self.selected_attempt_id = self._attempt_id_by_index[0]

    def _selected_attempt_id_from_listbox(self) -> str:
        sel = self.attempt_list.curselection()
        if not sel:
            return ""
        idx = sel[0]
        return self._attempt_id_by_index.get(idx, "")

    def view_attempt_details(self):
        eid = self.selected_exam_id or self._selected_exam_id()
        if not eid:
            err("Select an exam first.")
            return

        aid = self._selected_attempt_id_from_listbox()
        if not aid:
            err("Select a student attempt.")
            return

        attempts = self.app.store.list_attempts_for_exam(eid)
        target = next((x for x in attempts if x.attempt_id == aid), None)
        if not target:
            err("Attempt not found. Please reload.")
            return

        frame: TeacherAttemptFrame = self.app.frames["TeacherAttemptFrame"]  # type: ignore
        frame.load_attempt(target, back_to="TeacherFrame")
        self.app.show_frame("TeacherAttemptFrame")


    def export_selected_attempt_word(self):
        eid = self.selected_exam_id or self._selected_exam_id()
        if not eid:
            err("Select an exam first.")
            return

        aid = self._selected_attempt_id_from_listbox()
        if not aid:
            err("Select a student attempt first.")
            return

        e = self.app.store.get_exam(eid)
        if not e:
            err("Exam not found.")
            return

        attempts = self.app.store.list_attempts_for_exam(eid)
        a = next((x for x in attempts if x.attempt_id == aid), None)
        if not a:
            err("Attempt not found.")
            return

        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"attempt_{a.username}_{e.access_code}_{a.attempt_id}.docx"
        )
        if not filepath:
            return

        export_attempt_to_word(e, a, filepath)
        info("Exported attempt to Word.")

    def _on_attempt_click(self, event):
        idx = self.attempt_list.nearest(event.y)
        if idx >= 0:
            self.attempt_list.selection_clear(0, tk.END)
            self.attempt_list.selection_set(idx)
            self.attempt_list.activate(idx)
            self.attempt_list.focus_set()
            self.selected_attempt_id = self._attempt_id_by_index.get(idx, "")

    def _on_attempt_select(self, event=None):
        sel = self.attempt_list.curselection()
        if not sel:
            self.selected_attempt_id = ""
            return
        idx = sel[0]
        self.selected_attempt_id = self._attempt_id_by_index.get(idx, "")


# -----------------------------
# Student Frame
# -----------------------------
class StudentFrame(ttk.Frame):
    def __init__(self, parent, app: App):
        super().__init__(parent)
        self.app = app

        Header(self, "Student Panel", "Student must enter exam by CODE only.").pack(fill="x", pady=(0, 10))

        top = ttk.Frame(self)
        top.pack(fill="x")
        ttk.Button(top, text="Reload", command=self.app.reload_data).pack(side="right", padx=6)
        ttk.Button(top, text="Logout", command=self.app.logout).pack(side="right")

        main = ttk.Frame(self)
        main.pack(fill="both", expand=True, pady=10)

        left = ttk.LabelFrame(main, text="Enter exam by CODE")
        left.pack(side="left", fill="both", expand=True, padx=(0, 8))
        right = ttk.LabelFrame(main, text="Profile + My attempts (review may)")
        right.pack(side="left", fill="both", expand=True, padx=(8, 0))

        lf = ttk.Frame(left, padding=10)
        lf.pack(fill="x")
        ttk.Label(lf, text="CODE:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        self.code_var = tk.StringVar()
        ttk.Entry(lf, textvariable=self.code_var, width=18).grid(row=0, column=1, sticky="w", padx=6, pady=6)
        ttk.Button(lf, text="Open", command=self.open_by_code).grid(row=0, column=2, padx=6)

        ttk.Label(left, text="Note: You do not see exam list. Teacher should give you the code.").pack(anchor="w", padx=10)

        # profile
        pf = ttk.Frame(right, padding=10)
        pf.pack(fill="x")
        ttk.Label(pf, text="Full name:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(pf, text="DOB (YYYY-MM-DD):").grid(row=1, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(pf, text="Student ID:").grid(row=2, column=0, sticky="e", padx=6, pady=6)

        self.full_name = tk.StringVar()
        self.dob = tk.StringVar()
        self.sid = tk.StringVar()
        ttk.Entry(pf, textvariable=self.full_name, width=28).grid(row=0, column=1, sticky="w", padx=6, pady=6)
        ttk.Entry(pf, textvariable=self.dob, width=28).grid(row=1, column=1, sticky="w", padx=6, pady=6)
        ttk.Entry(pf, textvariable=self.sid, width=28).grid(row=2, column=1, sticky="w", padx=6, pady=6)
        ttk.Button(pf, text="Update profile", command=self.update_profile).grid(row=3, column=0, columnspan=2, pady=(8, 4))

        ttk.Separator(right).pack(fill="x", padx=10, pady=6)
        ttk.Label(right, text="My attempts:").pack(anchor="w", padx=10)
        self.attempt_list = tk.Listbox(right, height=16)
        self.attempt_list.pack(fill="both", expand=True, padx=10, pady=(6, 10))
        ttk.Button(right, text="Review selected attempt", command=self.review_selected).pack(pady=(0, 10))

    def on_show(self):
        if not self.app.current_user or self.app.current_user.role != "Student":
            err("You need Student role.")
            self.app.logout()
            return
        u = self.app.store.find_user(self.app.current_user.username)
        if u:
            self.app.current_user = u
        self.full_name.set(self.app.current_user.full_name)
        self.dob.set(self.app.current_user.dob)
        self.sid.set(self.app.current_user.student_id)
        self.refresh_attempts()

    def update_profile(self):
        ok = self.app.store.update_profile(
            self.app.current_user.username,
            self.full_name.get().strip(),
            self.dob.get().strip(),
            self.sid.get().strip()
        )
        if not ok:
            err("Update failed.")
            return
        u = self.app.store.find_user(self.app.current_user.username)
        if u:
            self.app.current_user = u
        info("Profile updated.")
        self.refresh_attempts()

    def refresh_attempts(self):
        self.attempt_list.delete(0, tk.END)
        attempts = self.app.store.list_attempts_for_user(self.app.current_user.username)
        if not attempts:
            self.attempt_list.insert(tk.END, "No attempts yet.")
            return
        for a in attempts:
            score10 = (a.score / max(1, a.total)) * 10.0
            self.attempt_list.insert(
                tk.END,
                f"{fmt_dt_full(a.submitted_at)} | {a.title} | code:{a.code} | {score10:.2f}/10"
            )

    def _check_password(self, exam: Exam) -> bool:
        if not exam.password:
            return True
        pw = simpledialog.askstring("Password", "This exam needs password:", show="*")
        if pw is None:
            return False
        if pw != exam.password:
            err("Password is not correct.")
            return False
        return True

    def open_by_code(self):
        code = self.code_var.get().strip().upper()
        if not code:
            err("Need code.")
            return
        exam = self.app.store.get_exam_by_code(code)
        if not exam:
            err("Exam not found by this code.")
            return

        now = int(time.time())
        if now < exam.start_ts:
            err(f"Exam not open yet.\nStart: {fmt_dt(exam.start_ts)}")
            return
        if now > exam.end_ts:
            err(f"Exam closed.\nEnd: {fmt_dt(exam.end_ts)}")
            return

        if exam.attempt_limit > 0:
            used = self.app.store.count_attempts_for_user_exam(self.app.current_user.username, exam.exam_id)
            if used >= exam.attempt_limit:
                err(f"Attempt limit reached.\nLimit: {exam.attempt_limit}, Used: {used}")
                return

        if not self._check_password(exam):
            return

        take: ExamTakeFrame = self.app.frames["ExamTakeFrame"]  # type: ignore
        take.load_exam(exam)
        self.app.show_frame("ExamTakeFrame")

    def review_selected(self):
        sel = self.attempt_list.curselection()
        if not sel:
            err("Select an attempt.")
            return
        attempts = self.app.store.list_attempts_for_user(self.app.current_user.username)
        idx = sel[0]
        if idx < 0 or idx >= len(attempts):
            err("Attempt not found.")
            return
        attempt = attempts[idx]
        exam = self.app.store.get_exam(attempt.exam_id)
        if not exam:
            err("Exam data not found.")
            return
        if not exam.allow_review:
            err("Teacher does not allow review.")
            return
        rv: ReviewFrame = self.app.frames["ReviewFrame"]  # type: ignore
        rv.load_review(exam, attempt, back_to="StudentFrame")
        self.app.show_frame("ReviewFrame")



# -----------------------------
# Exam Take Frame (Grid + Mark) - Đã sửa lỗi pady
# -----------------------------
class ExamTakeFrame(ttk.Frame):
    def __init__(self, parent, app: App):
        super().__init__(parent)
        self.app = app
        self.exam: Optional[Exam] = None
        self.index = 0
        self.answers: List[Set[int]] = []
        self.marked_questions: Set[int] = set()
        self.started_at: float = 0.0
        self.end_time: float = 0.0
        self._timer_job = None
        self._auto_submitted = False

        # Top bar
        top = ttk.Frame(self)
        top.pack(fill="x", side="top", pady=(0, 5))
        self.timer_label = ttk.Label(top, text="Time left: --:--", font=("Segoe UI", 12, "bold"), foreground="red")
        self.timer_label.pack(side="left", padx=10)
        ttk.Button(top, text="Exit", command=self.back).pack(side="right", padx=10)

        # Main Container (Split Left/Right)
        container = ttk.Frame(self)
        container.pack(fill="both", expand=True)

        # --- Sidebar (Right) ---
        self.sidebar = ttk.LabelFrame(container, text="Questions", padding=5)
        self.sidebar.pack(side="right", fill="y", padx=5, pady=5)
        
        self.grid_frame = ttk.Frame(self.sidebar)
        self.grid_frame.pack(fill="both", expand=True)
        
        legend = ttk.Frame(self.sidebar)
        legend.pack(fill="x", pady=10)
        tk.Label(legend, text="■ Current", fg="blue").pack(anchor="w")
        tk.Label(legend, text="■ Answered", fg="green").pack(anchor="w")
        tk.Label(legend, text="■ Marked", fg="orange").pack(anchor="w")

        # --- Content (Left) ---
        self.main_area = ttk.Frame(container, padding=10)
        self.main_area.pack(side="left", fill="both", expand=True)

        self.title_label = ttk.Label(self.main_area, text="", font=("Segoe UI", 14, "bold"))
        self.title_label.pack(anchor="w", pady=(0, 8))

        self.q_label = ttk.Label(self.main_area, text="", wraplength=700, justify="left", font=("Segoe UI", 11))
        self.q_label.pack(anchor="w", pady=(0, 10))

        self.opt_vars = [tk.IntVar(value=0) for _ in range(4)]
        self.check_buttons = []
        for i in range(4):
            cb = ttk.Checkbutton(self.main_area, text="", variable=self.opt_vars[i])
            cb.pack(anchor="w", pady=3)
            self.check_buttons.append(cb)

        # Bottom Nav (Đã sửa lỗi)
        nav = ttk.Frame(self.main_area) 
        nav.pack(fill="x", side="bottom", pady=20) # Chuyển pady xuống đây

        self.progress_label = ttk.Label(nav, text="")
        self.progress_label.pack(side="left")

        ttk.Button(nav, text="Submit Exam", command=self.submit).pack(side="right", padx=6)
        ttk.Button(nav, text="Next >", command=self.next_q).pack(side="right", padx=6)
        ttk.Button(nav, text="< Prev", command=self.prev_q).pack(side="right", padx=6)
        
        self.btn_mark = tk.Button(nav, text="Mark for Review", bg="lightyellow", command=self.toggle_mark)
        self.btn_mark.pack(side="right", padx=20)

        self.nav_buttons = []

    def load_exam(self, exam: Exam):
        self.stop_timer()
        self.exam = exam
        self.index = 0
        self.answers = [set() for _ in range(len(exam.questions))]
        self.marked_questions = set()
        self.started_at = time.time()
        self.end_time = self.started_at + exam.duration_seconds
        self._auto_submitted = False
        
        self.create_nav_grid()
        self.render()
        self._tick()

    def create_nav_grid(self):
        for w in self.grid_frame.winfo_children(): w.destroy()
        self.nav_buttons = []
        if not self.exam: return

        cols = 5
        for i in range(len(self.exam.questions)):
            # Use tk.Button for background colors
            btn = tk.Button(self.grid_frame, text=str(i + 1), width=4,
                            command=lambda idx=i: self.jump_to(idx))
            r = i // cols
            c = i % cols
            btn.grid(row=r, column=c, padx=2, pady=2)
            self.nav_buttons.append(btn)

    def jump_to(self, target):
        self._save_current()
        self.index = target
        self.render()

    def toggle_mark(self):
        if self.index in self.marked_questions:
            self.marked_questions.remove(self.index)
        else:
            self.marked_questions.add(self.index)
        self.render()

    def stop_timer(self):
        if self._timer_job:
            self.after_cancel(self._timer_job)
            self._timer_job = None

    def _tick(self):
        if not self.exam: return
        left = int(self.end_time - time.time())
        mm, ss = max(0, left) // 60, max(0, left) % 60
        self.timer_label.config(text=f"Time left: {mm:02d}:{ss:02d}")

        if left <= 0 and not self._auto_submitted:
            self._auto_submitted = True
            self._save_current()
            self._submit_internal(auto=True)
            return
        self._timer_job = self.after(1000, self._tick)

    def _current_selection(self) -> Set[int]:
        return {i for i in range(4) if int(self.opt_vars[i].get()) == 1}

    def _save_current(self):
        if self.exam: self.answers[self.index] = self._current_selection()

    def render(self):
        if not self.exam: return
        
        q = self.exam.questions[self.index]
        self.title_label.config(text=f"{self.exam.title}")
        self.q_label.config(text=f"Q{self.index+1}: {q.text}")
        
        for i in range(4): self.check_buttons[i].config(text=q.options[i])
        
        saved = self.answers[self.index]
        for i in range(4): self.opt_vars[i].set(1 if i in saved else 0)
            
        self.progress_label.config(text=f"Question: {self.index+1}/{len(self.exam.questions)}")

        # Mark button
        if self.index in self.marked_questions:
            self.btn_mark.config(text="Unmark Flag", bg="orange", fg="white")
        else:
            self.btn_mark.config(text="Mark for Review", bg="lightyellow", fg="black")

        # Color Grid
        for i, btn in enumerate(self.nav_buttons):
            bg, fg = "#f0f0f0", "black"
            
            if len(self.answers[i]) > 0: 
                bg = "#90ee90" # Green (Answered)
            
            if i in self.marked_questions: 
                bg = "orange" # Marked
                
            if i == self.index: 
                bg, fg = "blue", "white" # Current
                
            btn.config(bg=bg, fg=fg)

    def next_q(self):
        self._save_current()
        if self.index < len(self.exam.questions) - 1:
            self.index += 1
            self.render()
        else:
            info("Last question.")

    def prev_q(self):
        self._save_current()
        if self.index > 0:
            self.index -= 1
            self.render()

    def submit(self):
        self._save_current()
        done = sum(1 for a in self.answers if len(a) > 0)
        if not messagebox.askyesno("Submit", f"Answered: {done}/{len(self.exam.questions)}\nSubmit now?"):
            return
        self._submit_internal(False)

    def _submit_internal(self, auto):
        # Calculate Score
        total_score = 0.0
        for ans, q in zip(self.answers, self.exam.questions):
            earned, _, _ = score_question_partial(ans, set(q.correct_indices), 1.0)
            total_score += earned
        
        u = self.app.current_user
        if not u: return 

        a = Attempt(
            self.app.store.new_attempt_id(), self.exam.exam_id, self.exam.access_code, self.exam.title,
            u.username, u.full_name, u.student_id,
            total_score, len(self.exam.questions), self.started_at, time.time(),
            int(time.time() - self.started_at), [sorted(list(s)) for s in self.answers]
        )
        self.app.store.add_attempt(a)
        self.stop_timer()
        
        s10 = (total_score/max(1, len(self.exam.questions)))*10.0
        msg = f"Time is over. Auto submit.\nScore: {s10:.2f}/10" if auto else f"Submitted!\nScore: {s10:.2f}/10"
        messagebox.showinfo("Done", msg)
        self.back()

    def back(self):
        self.stop_timer()
        self.app.show_frame("StudentFrame")

# -----------------------------
# Review Frame
# -----------------------------
class ReviewFrame(ttk.Frame):
    def __init__(self, parent, app: App):
        super().__init__(parent)
        self.app = app
        self.exam: Optional[Exam] = None
        self.attempt: Optional[Attempt] = None
        self.back_to = "StudentFrame"

        Header(self, "Review", "Show answers, missing correct, extra wrong, and /10 score.").pack(fill="x", pady=(0, 10))

        top = ttk.Frame(self)
        top.pack(fill="x")
        self.info_label = ttk.Label(top, text="")
        self.info_label.pack(side="left")
        ttk.Button(top, text="Back", command=self.go_back).pack(side="right")

        self.text = tk.Text(self, wrap="word", height=32)
        self.text.pack(fill="both", expand=True)

    def load_review(self, exam: Exam, attempt: Attempt, back_to: str):
        self.exam = exam
        self.attempt = attempt
        self.back_to = back_to
        self.render()

    def render(self):
        self.text.config(state="normal")
        self.text.delete("1.0", tk.END)
        if not self.exam or not self.attempt:
            self.text.insert(tk.END, "No data.")
            self.text.config(state="disabled")
            return

        total_q = max(1, len(self.exam.questions))
        score10 = (self.attempt.score / total_q) * 10.0
        need_more = max(0.0, 10.0 - score10)

        took = f"{self.attempt.time_taken_seconds // 60:02d}:{self.attempt.time_taken_seconds % 60:02d}"
        submit_t = fmt_dt_full(self.attempt.submitted_at)
        self.info_label.config(text=f"{self.exam.title} | {self.attempt.username} | {score10:.2f}/10 | need {need_more:.2f} | took {took} | {submit_t}")

        points_per_q = 10.0 / total_q
        answers = self.attempt.answers or []

        self.text.insert(tk.END, f"Total: {score10:.2f}/10\nNeed: {need_more:.2f} more to reach 10\n")
        self.text.insert(tk.END, "Rule: partial credit for multi-correct questions.\n")

        for i, q in enumerate(self.exam.questions):
            user_sel = set(answers[i]) if i < len(answers) else set()
            correct = set(q.correct_indices)
            earned, missing, extra = score_question_partial(user_sel, correct, points=points_per_q)

            self.text.insert(tk.END, "\n" + "-" * 60 + "\n")
            self.text.insert(tk.END, f"Q{i+1}: {q.text}\n")
            self.text.insert(tk.END, f"Earned: {earned:.2f}/{points_per_q:.2f}\n")
            for oi, opt in enumerate(q.options):
                mu = "[x]" if oi in user_sel else "[ ]"
                mc = "(correct)" if oi in correct else ""
                self.text.insert(tk.END, f"  {mu} {oi+1}. {opt} {mc}\n")
            self.text.insert(tk.END, "Missing correct: " + (", ".join(str(x+1) for x in sorted(missing)) if missing else "none") + "\n")
            self.text.insert(tk.END, "Extra wrong: " + (", ".join(str(x+1) for x in sorted(extra)) if extra else "none") + "\n")

        self.text.config(state="disabled")

    def go_back(self):
        self.app.show_frame(self.back_to)


# -----------------------------
# Teacher Attempt Details
# -----------------------------
class TeacherAttemptFrame(ttk.Frame):
    def __init__(self, parent, app: App):
        super().__init__(parent)
        self.app = app
        self.attempt: Optional[Attempt] = None
        self.exam: Optional[Exam] = None
        self.back_to = "TeacherFrame"

        Header(self, "Attempt Details", "Teacher can see student answers with partial credit.").pack(fill="x", pady=(0, 10))

        top = ttk.Frame(self)
        top.pack(fill="x")
        self.info_label = ttk.Label(top, text="")
        self.info_label.pack(side="left")
        ttk.Button(top, text="Back", command=self.go_back).pack(side="right")

        self.text = tk.Text(self, wrap="word", height=32)
        self.text.pack(fill="both", expand=True)

    def load_attempt(self, attempt: Attempt, back_to: str):
        self.attempt = attempt
        self.back_to = back_to
        self.exam = self.app.store.get_exam(attempt.exam_id)
        self.render()

    def render(self):
        self.text.config(state="normal")
        self.text.delete("1.0", tk.END)
        if not self.attempt or not self.exam:
            self.text.insert(tk.END, "No data.")
            self.text.config(state="disabled")
            return

        total_q = max(1, len(self.exam.questions))
        score10 = (self.attempt.score / total_q) * 10.0
        submit_t = fmt_dt_full(self.attempt.submitted_at)
        name = self.attempt.full_name or "(no name)"
        self.info_label.config(text=f"{self.exam.title} | {name} | {self.attempt.username} | {score10:.2f}/10 | {submit_t}")

        points_per_q = 10.0 / total_q
        answers = self.attempt.answers or []

        for i, q in enumerate(self.exam.questions):
            user_sel = set(answers[i]) if i < len(answers) else set()
            correct = set(q.correct_indices)
            earned, missing, extra = score_question_partial(user_sel, correct, points=points_per_q)

            self.text.insert(tk.END, "\n" + "-" * 60 + "\n")
            self.text.insert(tk.END, f"Q{i+1}: {q.text}\n")
            self.text.insert(tk.END, f"Earned: {earned:.2f}/{points_per_q:.2f}\n")
            for oi, opt in enumerate(q.options):
                mu = "[x]" if oi in user_sel else "[ ]"
                mc = "(correct)" if oi in correct else ""
                self.text.insert(tk.END, f"  {mu} {oi+1}. {opt} {mc}\n")
            self.text.insert(tk.END, "Missing correct: " + (", ".join(str(x+1) for x in sorted(missing)) if missing else "none") + "\n")
            self.text.insert(tk.END, "Extra wrong: " + (", ".join(str(x+1) for x in sorted(extra)) if extra else "none") + "\n")

        self.text.config(state="disabled")

    def go_back(self):
        self.app.show_frame(self.back_to)


# -----------------------------
# Template Preview
# -----------------------------
class TemplatePreviewFrame(ttk.Frame):
    def __init__(self, parent, app: App):
        super().__init__(parent)
        self.app = app
        self.template_id: Optional[str] = None
        self.back_to = "TeacherFrame"

        Header(self, "Template Preview", "Template is in warehouse (no code).").pack(fill="x", pady=(0, 10))

        top = ttk.Frame(self)
        top.pack(fill="x")
        self.info_label = ttk.Label(top, text="")
        self.info_label.pack(side="left")
        ttk.Button(top, text="Back", command=self.go_back).pack(side="right")
        ttk.Button(top, text="Reload", command=self.app.reload_data).pack(side="right", padx=6)

        self.text = tk.Text(self, wrap="word", height=32)
        self.text.pack(fill="both", expand=True)

    def load_template(self, template_id: str, back_to: str):
        self.template_id = template_id
        self.back_to = back_to
        self.render()

    def on_show(self):
        if self.template_id:
            self.render()

    def render(self):
        self.text.config(state="normal")
        self.text.delete("1.0", tk.END)
        if not self.template_id:
            self.text.insert(tk.END, "No template.")
            self.text.config(state="disabled")
            return
        t = self.app.store.get_template(self.template_id)
        if not t:
            self.text.insert(tk.END, "Template not found.")
            self.text.config(state="disabled")
            return
        self.info_label.config(text=f"{t.title} | id:{t.template_id} | questions:{len(t.questions)}")
        for i, q in enumerate(t.questions):
            correct = ", ".join(str(x+1) for x in sorted(q.correct_indices))
            self.text.insert(tk.END, "\n" + "-" * 60 + "\n")
            self.text.insert(tk.END, f"Q{i+1}: {q.text}\n")
            for oi, opt in enumerate(q.options):
                mark = "(correct)" if oi in set(q.correct_indices) else ""
                self.text.insert(tk.END, f"  {oi+1}. {opt} {mark}\n")
            self.text.insert(tk.END, f"Correct: {correct}\n")
        self.text.config(state="disabled")

    def go_back(self):
        self.app.show_frame(self.back_to)


# -----------------------------
# Exam Preview
# -----------------------------
class ExamPreviewFrame(ttk.Frame):
    def __init__(self, parent, app: App):
        super().__init__(parent)
        self.app = app
        self.exam_id: Optional[str] = None
        self.back_to = "TeacherFrame"

        Header(self, "Exam Preview", "Exam has code + time window.").pack(fill="x", pady=(0, 10))

        top = ttk.Frame(self)
        top.pack(fill="x")
        self.info_label = ttk.Label(top, text="")
        self.info_label.pack(side="left")
        ttk.Button(top, text="Back", command=self.go_back).pack(side="right")
        ttk.Button(top, text="Reload", command=self.app.reload_data).pack(side="right", padx=6)

        self.text = tk.Text(self, wrap="word", height=32)
        self.text.pack(fill="both", expand=True)

    def load_exam(self, exam_id: str, back_to: str):
        self.exam_id = exam_id
        self.back_to = back_to
        self.render()

    def on_show(self):
        if self.exam_id:
            self.render()

    def render(self):
        self.text.config(state="normal")
        self.text.delete("1.0", tk.END)
        if not self.exam_id:
            self.text.insert(tk.END, "No exam.")
            self.text.config(state="disabled")
            return
        e = self.app.store.get_exam(self.exam_id)
        if not e:
            self.text.insert(tk.END, "Exam not found.")
            self.text.config(state="disabled")
            return

        self.info_label.config(text=f"{e.title} | id:{e.exam_id} | code:{e.access_code}")
        self.text.insert(tk.END, f"Title: {e.title}\n")
        self.text.insert(tk.END, f"Code: {e.access_code}\n")
        self.text.insert(tk.END, f"Window: {fmt_dt(e.start_ts)} -> {fmt_dt(e.end_ts)}\n")
        self.text.insert(tk.END, f"Duration: {max(1, e.duration_seconds//60)} minutes\n")
        self.text.insert(tk.END, f"Password: {'set' if e.password else 'none'}\n")
        self.text.insert(tk.END, f"Allow review: {'yes' if e.allow_review else 'no'}\n")
        self.text.insert(tk.END, f"Attempt limit: {e.attempt_limit} (0=unlimited)\n")
        self.text.insert(tk.END, f"From template: {e.template_id}\n")

        for i, q in enumerate(e.questions):
            correct = ", ".join(str(x+1) for x in sorted(q.correct_indices))
            self.text.insert(tk.END, "\n" + "-" * 60 + "\n")
            self.text.insert(tk.END, f"Q{i+1}: {q.text}\n")
            for oi, opt in enumerate(q.options):
                mark = "(correct)" if oi in set(q.correct_indices) else ""
                self.text.insert(tk.END, f"  {oi+1}. {opt} {mark}\n")
            self.text.insert(tk.END, f"Correct: {correct}\n")

        self.text.config(state="disabled")

    def go_back(self):
        self.app.show_frame(self.back_to)


# -----------------------------
# Main
# -----------------------------
def main():
    app = App()
    app.mainloop()


if __name__ == "__main__":
    main()