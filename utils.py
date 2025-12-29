# utils.py
import time
from tkinter import messagebox
from typing import List, Optional, Set, Tuple

from models import Template, Exam, Attempt, score_question_partial

try:
    from docx import Document
    from openpyxl import Workbook
except ImportError:
    pass

# -----------------------------
# Time helpers
# -----------------------------
def parse_dt(s: str) -> Optional[int]:
    """Parse 'YYYY-MM-DD HH:MM' to unix timestamp (local)."""
    s = (s or "").strip()
    try:
        return int(time.mktime(time.strptime(s, "%Y-%m-%d %H:%M")))
    except Exception:
        return None

def fmt_dt(ts: int) -> str:
    try:
        return time.strftime("%Y-%m-%d %H:%M", time.localtime(int(ts)))
    except Exception:
        return "N/A"

def fmt_dt_full(ts: float) -> str:
    try:
        return time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(float(ts)))
    except Exception:
        return "N/A"

# -----------------------------
# Export helpers
# -----------------------------
def export_template_to_word(t: Template, filepath: str, include_answers: bool = True):
    try:
        doc = Document()
        doc.add_heading(f"Template: {t.title}", level=1)
        doc.add_paragraph(f"Template ID: {t.template_id}")
        doc.add_paragraph(f"Created by: {t.created_by}")
        doc.add_paragraph(f"Questions: {len(t.questions)}")

        for i, q in enumerate(t.questions, start=1):
            doc.add_heading(f"Q{i}: {q.text}", level=2)
            for oi, opt in enumerate(q.options, start=1):
                doc.add_paragraph(f"{oi}. {opt}", style="List Bullet")
            if include_answers:
                correct = ", ".join(str(x + 1) for x in sorted(q.correct_indices))
                doc.add_paragraph(f"Correct: {correct}")
        doc.save(filepath)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export docx: {e}")

def export_exam_results_to_excel(e: Exam, attempts: List[Attempt], filepath: str):
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Results"

        ws.append([
            "Exam Title", "Exam ID", "Code",
            "Username", "Full Name", "Student ID",
            "Score (/10)", "Raw Score", "Total Questions",
            "Time Taken (sec)", "Submitted At"
        ])

        for a in attempts:
            score10 = (a.score / max(1, a.total)) * 10.0
            ws.append([
                e.title, e.exam_id, e.access_code,
                a.username, a.full_name, a.student_id,
                round(score10, 2), round(a.score, 4), a.total,
                a.time_taken_seconds, fmt_dt_full(a.submitted_at)
            ])
        wb.save(filepath)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export xlsx: {e}")