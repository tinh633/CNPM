
from __future__ import annotations

import json
import os
import getpass
import time
import datetime
import secrets
import string
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Set, Any, Tuple

DATA_FILE = "quiz_data.json"
ROLES = ["Admin", "Teacher", "Student"]
ROLE_CANON = {"admin": "Admin", "teacher": "Teacher", "student": "Student"}

# Optional export libs
try:
    from docx import Document  # type: ignore
except Exception:
    Document = None  # type: ignore

try:
    from openpyxl import Workbook  # type: ignore
except Exception:
    Workbook = None  # type: ignore


# -----------------------------
# Date/Time helpers
# -----------------------------
def parse_dt(s: str) -> Optional[int]:
    """Parse 'YYYY-MM-DD HH:MM' to unix timestamp (local). Returns None if invalid."""
    s = (s or "").strip()
    try:
        return int(time.mktime(time.strptime(s, "%Y-%m-%d %H:%M")))
    except Exception:
        return None


def fmt_dt(ts: int) -> str:
    try:
        return time.strftime("%Y-%m-%d %H:%M", time.localtime(int(ts)))
    except Exception:
        return "N/A"


def fmt_dt_full(ts: float) -> str:
    try:
        return time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(float(ts)))
    except Exception:
        return "N/A"


def is_valid_date_yyyy_mm_dd(s: str) -> bool:
    """Strict check: YYYY-MM-DD with month 1-12, day valid for month/leap year."""
    s = (s or "").strip()
    try:
        y, m, d = s.split("-")
        y = int(y)
        m = int(m)
        d = int(d)
        if y < 1900 or y > 2100:
            return False
        if m < 1 or m > 12:
            return False
        # days in month
        dim = [31, 29 if _is_leap(y) else 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31][m - 1]
        if d < 1 or d > dim:
            return False
        return True
    except Exception:
        return False


def _is_leap(y: int) -> bool:
    return (y % 4 == 0 and y % 100 != 0) or (y % 400 == 0)


def pause():
    input("\nPress ENTER to continue...")


def audit_log(action: str, performed_by: str, target: str = "", reason: str = ""):
    """Append a simple audit entry to audit.log (UTF-8)."""
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{ts}] action={action} by={performed_by}"
    if target:
        line += f" target={target}"
    if reason:
        safe_reason = " ".join((reason or "").split())
        line += f" reason={safe_reason}"
    line += "\n"
    try:
        with open("audit.log", "a", encoding="utf-8") as f:
            f.write(line)
    except Exception:
        pass


# -----------------------------
# Data Models
# -----------------------------
@dataclass
class User:
    username: str
    password: str
    role: str
    must_change_password: bool = False
    full_name: str = ""
    dob: str = ""         # YYYY-MM-DD
    student_id: str = ""  # for students


@dataclass
class Question:
    text: str
    options: List[str]              # 4 options
    correct_indices: List[int]      # can be multiple


@dataclass
class Template:
    template_id: str
    title: str
    created_by: str
    questions: List[Question]


@dataclass
class Exam:
    exam_id: str
    template_id: str
    title: str
    created_by: str
    access_code: str
    password: str
    duration_seconds: int
    allow_review: bool
    attempt_limit: int             # 0 means unlimited
    start_ts: int
    end_ts: int
    questions: List[Question]


@dataclass
class Attempt:
    attempt_id: str
    exam_id: str
    code: str
    title: str
    username: str
    full_name: str
    student_id: str
    score: float                   # max = total questions (each question max 1)
    total: int
    started_at: float
    submitted_at: float
    time_taken_seconds: int
    answers: List[List[int]]


# -----------------------------
# Scoring (partial credit)
# -----------------------------
def score_question_partial(user_sel: Set[int], correct: Set[int], points: float) -> Tuple[float, Set[int], Set[int]]:
    """
    points: points for this question
    rule:
      earned_ratio = max(0, (c - w) / k)
      k = number of correct options
      c = selected correct count
      w = selected wrong count
    """
    if not correct:
        return (0.0, set(), set())
    c = len(user_sel & correct)
    w = len(user_sel - correct)
    k = len(correct)
    earned_ratio = (c - w) / k
    earned_ratio = max(0.0, min(1.0, earned_ratio))
    earned = earned_ratio * points
    missing = correct - user_sel
    extra = user_sel - correct
    return earned, missing, extra


# -----------------------------
# Storage + migrations
# -----------------------------
class DataStore:
    def __init__(self, path: str):
        self.path = path
        self.data: Dict[str, Any] = {"users": [], "templates": [], "exams": [], "attempts": []}
        self.load()

    def load(self):
        if not os.path.exists(self.path):
            self._seed_default()
            self.save()
            return

        try:
            with open(self.path, "r", encoding="utf-8") as f:
                raw = json.load(f)
        except Exception:
            self._seed_default()
            self.save()
            return

        self.data = raw if isinstance(raw, dict) else {}
        self.data.setdefault("users", [])
        self.data.setdefault("templates", [])
        self.data.setdefault("exams", [])
        self.data.setdefault("attempts", [])

        # Normalize users
        for u in self.data["users"]:
            u.setdefault("full_name", "")
            u.setdefault("dob", "")
            u.setdefault("student_id", "")
            role = str(u.get("role", "")).strip()
            role_norm = ROLE_CANON.get(role.lower(), role)
            if role_norm not in ROLES:
                role_norm = "Student"
            u["role"] = role_norm

        # migrate templates/exams questions
        for t in self.data["templates"]:
            t.setdefault("questions", [])
            for qu in t.get("questions", []):
                if "correct_indices" not in qu and "correct_index" in qu:
                    qu["correct_indices"] = [int(qu.get("correct_index", 0))]
                qu.pop("correct_index", None)

        for e in self.data["exams"]:
            e.setdefault("attempt_limit", 0)
            e.setdefault("allow_review", False)
            e.setdefault("password", "")
            e.setdefault("duration_seconds", 0)
            e.setdefault("start_ts", 0)
            e.setdefault("end_ts", 0)
            e.setdefault("questions", [])
            for qu in e.get("questions", []):
                if "correct_indices" not in qu and "correct_index" in qu:
                    qu["correct_indices"] = [int(qu.get("correct_index", 0))]
                qu.pop("correct_index", None)

        for a in self.data["attempts"]:
            try:
                a["score"] = float(a.get("score", 0))
            except Exception:
                a["score"] = 0.0
            a.setdefault("answers", [])

        self.save()

    def save(self):
        with open(self.path, "w", encoding="utf-8") as f:
            json.dump(self.data, f, indent=2, ensure_ascii=False)

    def reset_to_default(self):
        """Reset all data to default demo users (admin/teacher/student)."""
        self._seed_default()
        self.save()

    def _seed_default(self):
        self.data = {
            "users": [
                asdict(User(username="default_admin", password="123@System", role="Admin")),
            ],
            "templates": [],
            "exams": [],
            "attempts": [],
        }

    # ---- Users ----
    def find_user(self, username: str) -> Optional[User]:
        for u in self.data["users"]:
            if u.get("username") == username:
                return User(**u)
        return None

    def list_users(self) -> List[User]:
        return [User(**u) for u in self.data["users"]]

    def add_user(self, user: User) -> bool:
        if self.find_user(user.username):
            return False
        user.role = ROLE_CANON.get(user.role.lower(), user.role)
        if user.role not in ROLES:
            user.role = "Student"
        self.data["users"].append(asdict(user))
        self.save()
        return True
    
    def delete_user(self, username: str) -> bool:
        """Delete a user by username (case-insensitive). Returns True if deleted."""
        uname = (username or "").strip().lower()
        if not uname:
            return False
        before = len(self.data.get("users", []))
        self.data["users"] = [u for u in self.data.get("users", []) if (u.get("username") or "").lower() != uname]
        after = len(self.data.get("users", []))
        if after == before:
            return False
        self.save()
        return True

    def update_password(self, username: str, new_password: str) -> bool:
        for u in self.data["users"]:
            if u.get("username") == username:
                u["password"] = new_password
                self.save()
                return True
        return False

    def set_must_change_password(self, username: str, value: bool) -> bool:
        for u in self.data.get("users", []):
            if (u.get("username") or "").lower() == (username or "").lower():
                u["must_change_password"] = bool(value)
                self.save()
                return True
        return False

    def update_profile(self, username: str, full_name: str, dob: str, student_id: str) -> bool:
        for u in self.data["users"]:
            if u.get("username") == username:
                u["full_name"] = full_name
                u["dob"] = dob
                u["student_id"] = student_id
                self.save()
                return True
        return False

    # ---- IDs/Codes ----
    def new_template_id(self) -> str:
        return f"TP{int(time.time() * 1000)}"

    def new_exam_id(self) -> str:
        return f"EX{int(time.time() * 1000)}"

    def new_attempt_id(self) -> str:
        return f"AT{int(time.time() * 1000)}{secrets.randbelow(1000)}"

    def _all_codes(self) -> Set[str]:
        return {(e.get("access_code") or "").upper() for e in self.data["exams"] if e.get("access_code")}

    def new_unique_code(self, length: int = 8) -> str:
        alphabet = string.ascii_uppercase + string.digits
        existing = self._all_codes()
        while True:
            code = "".join(secrets.choice(alphabet) for _ in range(length))
            if code not in existing:
                return code

    # ---- Templates ----
    def add_template(self, t: Template):
        self.data["templates"].append(self._template_to_dict(t))
        self.save()

    def list_templates(self) -> List[Template]:
        return [self._dict_to_template(x) for x in self.data["templates"]]

    def list_templates_by_teacher(self, teacher: str) -> List[Template]:
        return [t for t in self.list_templates() if t.created_by == teacher]

    def get_template(self, template_id: str) -> Optional[Template]:
        for t in self.data["templates"]:
            if t.get("template_id") == template_id:
                return self._dict_to_template(t)
        return None

    def delete_template(self, template_id: str) -> bool:
        before = len(self.data["templates"])
        self.data["templates"] = [t for t in self.data["templates"] if t.get("template_id") != template_id]
        after = len(self.data["templates"])
        if after == before:
            return False
        self.save()
        return True

    def has_exam_from_template(self, template_id: str) -> bool:
        return any(e.get("template_id") == template_id for e in self.data["exams"])

    def _template_to_dict(self, t: Template) -> Dict[str, Any]:
        return {
            "template_id": t.template_id,
            "title": t.title,
            "created_by": t.created_by,
            "questions": [asdict(q) for q in t.questions],
        }

    def _dict_to_template(self, d: Dict[str, Any]) -> Template:
        qs = []
        for q in d.get("questions", []):
            qs.append(Question(text=q["text"], options=q["options"], correct_indices=q.get("correct_indices", [])))
        return Template(
            template_id=d.get("template_id", ""),
            title=d.get("title", ""),
            created_by=d.get("created_by", ""),
            questions=qs
        )

    # ---- Exams ----
    def add_exam(self, e: Exam):
        self.data["exams"].append(self._exam_to_dict(e))
        self.save()

    def list_exams(self) -> List[Exam]:
        return [self._dict_to_exam(x) for x in self.data["exams"]]

    def list_exams_by_teacher(self, teacher: str) -> List[Exam]:
        return [e for e in self.list_exams() if e.created_by == teacher]

    def get_exam_by_code(self, code: str) -> Optional[Exam]:
        code = (code or "").strip().upper()
        for e in self.data["exams"]:
            if (e.get("access_code") or "").upper() == code:
                return self._dict_to_exam(e)
        return None

    def get_exam(self, exam_id: str) -> Optional[Exam]:
        for e in self.data["exams"]:
            if e.get("exam_id") == exam_id:
                return self._dict_to_exam(e)
        return None

    def delete_exam(self, exam_id: str) -> bool:
        before = len(self.data["exams"])
        self.data["exams"] = [e for e in self.data["exams"] if e.get("exam_id") != exam_id]
        after = len(self.data["exams"])
        if after == before:
            return False
        self.save()
        return True

    def _exam_to_dict(self, e: Exam) -> Dict[str, Any]:
        return {
            "exam_id": e.exam_id,
            "template_id": e.template_id,
            "title": e.title,
            "created_by": e.created_by,
            "access_code": e.access_code,
            "password": e.password,
            "duration_seconds": e.duration_seconds,
            "allow_review": e.allow_review,
            "attempt_limit": e.attempt_limit,
            "start_ts": e.start_ts,
            "end_ts": e.end_ts,
            "questions": [asdict(q) for q in e.questions],
        }

    def _dict_to_exam(self, d: Dict[str, Any]) -> Exam:
        qs = []
        for q in d.get("questions", []):
            qs.append(Question(text=q["text"], options=q["options"], correct_indices=q.get("correct_indices", [])))
        return Exam(
            exam_id=d.get("exam_id", ""),
            template_id=d.get("template_id", ""),
            title=d.get("title", ""),
            created_by=d.get("created_by", ""),
            access_code=d.get("access_code", ""),
            password=d.get("password", "") or "",
            duration_seconds=int(d.get("duration_seconds", 0) or 0),
            allow_review=bool(d.get("allow_review", False)),
            attempt_limit=int(d.get("attempt_limit", 0) or 0),
            start_ts=int(d.get("start_ts", 0) or 0),
            end_ts=int(d.get("end_ts", 0) or 0),
            questions=qs
        )

    # ---- Attempts ----
    def add_attempt(self, a: Attempt):
        self.data["attempts"].append(asdict(a))
        self.save()

    def list_attempts_for_user(self, username: str) -> List[Attempt]:
        out = [Attempt(**x) for x in self.data["attempts"] if x.get("username") == username]
        out.sort(key=lambda z: z.submitted_at, reverse=True)
        return out

    def list_attempts_for_exam(self, exam_id: str) -> List[Attempt]:
        out = [Attempt(**x) for x in self.data["attempts"] if x.get("exam_id") == exam_id]
        out.sort(key=lambda z: z.submitted_at, reverse=True)
        return out

    def count_attempts_for_user_exam(self, username: str, exam_id: str) -> int:
        return sum(1 for x in self.data["attempts"] if x.get("username") == username and x.get("exam_id") == exam_id)

    def delete_attempts_for_exam(self, exam_id: str) -> int:
        before = len(self.data["attempts"])
        self.data["attempts"] = [a for a in self.data["attempts"] if a.get("exam_id") != exam_id]
        after = len(self.data["attempts"])
        deleted = before - after
        if deleted:
            self.save()
        return deleted


# -----------------------------
# Export helpers
# -----------------------------
def export_template_to_word(t: Template, filepath: str, include_answers: bool = True):
    if Document is None:
        print("Export needs python-docx. Run: pip install python-docx")
        return
    doc = Document()
    doc.add_heading(f"Template: {t.title}", level=1)
    doc.add_paragraph(f"Template ID: {t.template_id}")
    doc.add_paragraph(f"Created by: {t.created_by}")
    doc.add_paragraph(f"Questions: {len(t.questions)}")
    for i, q in enumerate(t.questions, start=1):
        doc.add_heading(f"Q{i}: {q.text}", level=2)
        for oi, opt in enumerate(q.options, start=1):
            doc.add_paragraph(f"{oi}. {opt}", style="List Bullet")
        if include_answers:
            correct = ", ".join(str(x + 1) for x in sorted(q.correct_indices))
            doc.add_paragraph(f"Correct: {correct}")
    doc.save(filepath)
    print(f"Saved: {filepath}")


def export_exam_to_word(e: Exam, filepath: str, include_answers: bool = True):
    if Document is None:
        print("Export needs python-docx. Run: pip install python-docx")
        return
    doc = Document()
    doc.add_heading(f"Exam: {e.title}", level=1)
    doc.add_paragraph(f"Exam ID: {e.exam_id}")
    doc.add_paragraph(f"Code: {e.access_code}")
    doc.add_paragraph(f"Window: {fmt_dt(e.start_ts)} -> {fmt_dt(e.end_ts)}")
    doc.add_paragraph(f"Duration: {max(1, e.duration_seconds // 60)} minutes")
    doc.add_paragraph(f"Password: {'set' if e.password else 'none'}")
    doc.add_paragraph(f"Allow review: {'yes' if e.allow_review else 'no'}")
    doc.add_paragraph(f"Attempt limit: {e.attempt_limit} (0=unlimited)")
    doc.add_paragraph(f"From template: {e.template_id}")
    for i, q in enumerate(e.questions, start=1):
        doc.add_heading(f"Q{i}: {q.text}", level=2)
        for oi, opt in enumerate(q.options, start=1):
            doc.add_paragraph(f"{oi}. {opt}", style="List Bullet")
        if include_answers:
            correct = ", ".join(str(x + 1) for x in sorted(q.correct_indices))
            doc.add_paragraph(f"Correct: {correct}")
    doc.save(filepath)
    print(f"Saved: {filepath}")


def export_exam_results_to_excel(e: Exam, attempts: List[Attempt], filepath: str):
    if Workbook is None:
        print("Export needs openpyxl. Run: pip install openpyxl")
        return
    wb = Workbook()
    ws = wb.active
    ws.title = "Results"
    ws.append([
        "Exam Title", "Exam ID", "Code",
        "Username", "Full Name", "Student ID",
        "Score (/10)", "Raw Score", "Total Questions",
        "Time Taken (sec)", "Submitted At"
    ])
    total_q = max(1, len(e.questions))
    for a in attempts:
        score10 = (a.score / max(1, a.total)) * 10.0
        ws.append([
            e.title, e.exam_id, e.access_code,
            a.username, a.full_name, a.student_id,
            round(score10, 2), round(a.score, 4), a.total,
            a.time_taken_seconds, fmt_dt_full(a.submitted_at)
        ])
    wb.save(filepath)
    print(f"Saved: {filepath}")


def export_attempt_to_word(e: Exam, a: Attempt, filepath: str):
    if Document is None:
        print("Export needs python-docx. Run: pip install python-docx")
        return
    doc = Document()
    doc.add_heading("Attempt Report", level=1)
    doc.add_paragraph(f"Exam: {e.title}")
    doc.add_paragraph(f"Exam ID: {e.exam_id} | Code: {e.access_code}")
    doc.add_paragraph(f"Student: {a.full_name} | Username: {a.username} | Student ID: {a.student_id}")
    doc.add_paragraph(f"Started: {fmt_dt_full(a.started_at)}")
    doc.add_paragraph(f"Submitted: {fmt_dt_full(a.submitted_at)}")
    doc.add_paragraph(f"Time taken: {a.time_taken_seconds} sec")

    total_q = max(1, len(e.questions))
    score10 = (a.score / total_q) * 10.0
    doc.add_paragraph(f"Score: {score10:.2f}/10 (raw {a.score:.4f}/{total_q})")

    answers = a.answers or []
    points_per_q = 10.0 / total_q
    for i, q in enumerate(e.questions):
        user_sel = set(answers[i]) if i < len(answers) else set()
        correct = set(q.correct_indices)
        earned, missing, extra = score_question_partial(user_sel, correct, points=points_per_q)
        doc.add_heading(f"Q{i+1}: {q.text}", level=2)
        doc.add_paragraph(f"Earned: {earned:.2f}/{points_per_q:.2f}")
        for oi, opt in enumerate(q.options, start=1):
            mu = "[x]" if (oi - 1) in user_sel else "[ ]"
            mc = "(correct)" if (oi - 1) in correct else ""
            doc.add_paragraph(f"{mu} {oi}. {opt} {mc}")
        doc.add_paragraph("Missing correct: " + (", ".join(str(x+1) for x in sorted(missing)) if missing else "none"))
        doc.add_paragraph("Extra wrong: " + (", ".join(str(x+1) for x in sorted(extra)) if extra else "none"))
    doc.save(filepath)
    print(f"Saved: {filepath}")


# -----------------------------
# Console UI helpers
# -----------------------------
def clear_screen():
    os.system("cls" if os.name == "nt" else "clear")


def ask(prompt: str) -> str:
    return input(prompt).strip()


def ask_password(prompt: str) -> str:
    return input(prompt).strip()


def generate_strong_password(length: int = 12) -> str:
    """Generate a strong temporary password (>=8, upper+lower+digit+special)."""
    if length < 8:
        length = 8
    upper = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    lower = "abcdefghijklmnopqrstuvwxyz"
    digits = "0123456789"
    special = "!@#$%^&*()-_=+[]{};:,.?/"
    pwd = [secrets.choice(upper), secrets.choice(lower), secrets.choice(digits), secrets.choice(special)]
    all_chars = upper + lower + digits + special
    while len(pwd) < length:
        pwd.append(secrets.choice(all_chars))
    secrets.SystemRandom().shuffle(pwd)
    return "".join(pwd)


def validate_password_policy(pw: str, min_len: int = 8):
    """Return (ok, message). Policy: >=min_len, upper+lower+digit+special."""
    pw = (pw or "").strip()
    if len(pw) < min_len:
        return False, f"Password must be at least {min_len} characters."
    if not any(c.islower() for c in pw):
        return False, "Password must contain at least one lowercase letter."
    if not any(c.isupper() for c in pw):
        return False, "Password must contain at least one uppercase letter."
    if not any(c.isdigit() for c in pw):
        return False, "Password must contain at least one digit."
    if not any(not c.isalnum() for c in pw):
        return False, "Password must contain at least one special character."
    return True, ""

def ask_non_empty(prompt: str) -> str:
    while True:
        s = ask(prompt)
        if s:
            return s
        print("Please enter a value.")


def validate_username_12_chars(username: str) -> bool:
    return len(username) == 12 and username.isdigit()


def ask_int(prompt: str, min_v: int, max_v: int) -> int:
    while True:
        s = ask(prompt)
        try:
            v = int(s)
        except Exception:
            print("Please enter an integer.")
            continue
        if v < min_v or v > max_v:
            print(f"Range: {min_v}..{max_v}")
            continue
        return v


def ask_yes_no(prompt: str, default: bool = False) -> bool:
    suffix = " [Y/n] " if default else " [y/N] "
    s = ask(prompt + suffix).lower()
    if not s:
        return default
    return s in ("y", "yes")


def ask_dt(prompt: str) -> int:
    while True:
        s = ask(prompt + " (YYYY-MM-DD HH:MM): ")
        ts = parse_dt(s)
        if ts is None:
            print("Invalid datetime format.")
            continue
        return ts


def ask_dob(prompt: str = "DOB") -> str:
    while True:
        s = ask(prompt + " (YYYY-MM-DD, blank to skip): ")
        if not s:
            return ""
        if is_valid_date_yyyy_mm_dd(s):
            return s
        print("Invalid date. Month must be 1..12, day must be valid for the month.")


def ask_multi_choice(max_opt: int = 4) -> Set[int]:
    """
    Return 0-based indices.
    Input examples:
      2
      1 3
      1,3,4
      blank -> empty
    """
    s = ask("Your answer (e.g. 2 or 1 3). Blank to skip: ").strip()
    if not s:
        return set()
    parts = s.replace(",", " ").split()
    out: Set[int] = set()
    for p in parts:
        try:
            k = int(p)
        except Exception:
            continue
        if 1 <= k <= max_opt:
            out.add(k - 1)
    return out


def print_header(title: str, subtitle: str = ""):
    print("=" * 70)
    print(title)
    if subtitle:
        print(subtitle)
    print("=" * 70)


# -----------------------------
# Flows
# -----------------------------
def login_flow(store: DataStore) -> Optional[User]:
    clear_screen()
    print_header("Quiz Examination System (Console)", "Default System Admin: default_admin / 123@System")
    print("1) Login")
    print("0) Exit")
    c = ask("Choose: ")
    if c == "0":
        return User(username="__EXIT__", password="", role="")
    if c != "1":
        return None

    role = ask_non_empty("Role (Admin/Teacher/Student): ")
    role = ROLE_CANON.get(role.lower(), role)
    if role not in ROLES:
        print("Invalid role.")
        pause()
        return None

    username = ask_non_empty("Username: ")
    password = ask_non_empty("Password: ")

    u = store.find_user(username)
    if not u:
        print("User not found.")
        pause()
        return None
    if u.password != password:
        print("Password is not correct.")
        pause()
        return None
    u.role = ROLE_CANON.get(u.role.lower(), u.role)
    if u.role != role:
        print(f"Role mismatch. Account role: {u.role}. You selected: {role}")
        pause()
        return None

    display_name = (u.full_name or u.username)
    print(f"Login successful! Welcome {display_name} ({u.role}).")

    # Force password change at first login if required
    if getattr(u, "must_change_password", False):
        print("")
        print("⚠️ You must change your password at first login.")
        change_password_console(store, u)
        u2 = store.find_user(u.username)
        if u2 and getattr(u2, "must_change_password", False):
            print("Password change required. Login canceled.")
            pause()
            return None
        u = u2 or u

    pause()
    return u


# -----------------------------
# Admin
# -----------------------------
def admin_menu(store: DataStore, me: User):
    while True:
        clear_screen()
        print_header("Admin Panel", "Admin manages users.")
        print("1) List users")
        print("2) Create user")
        print("3) Reset user password")
        print("4) Delete user")
        print("0) Logout")
        c = ask("Choose: ")

        if c == "0":
            return
        if c == "1":
            clear_screen()
            print_header("Users")
            for u in store.list_users():
                print(f"- {u.username:12} | {u.role:7} | {u.full_name:20} | {u.dob:10} | {u.student_id}")
            pause()
        elif c == "2":
            clear_screen()
            print_header("Create User", "Admin enters only username and role. System generates a temporary password.")
            username = ask_non_empty("New username: ")
            if not validate_username_12_chars(username):
                print("❌ Username must be exactly 12 digits.")
                pause()
                continue
            role = ask_non_empty("Role (Admin/Teacher/Student): ")
            role = ROLE_CANON.get(role.lower(), role)
            if role not in ROLES:
                print("Invalid role.")
                pause()
                continue
            if store.find_user(username):
                print("Username exists.")
                pause()
                continue

            temp_pw = generate_strong_password(12)
            ok = store.add_user(User(username=username, password=temp_pw, role=role, must_change_password=True))
            if ok:
                audit_log("CREATE_ACCOUNT", performed_by=me.username, target=username, reason=f"role={role}")
                print("✅ Account created successfully!")
                print(f"Temporary password (show once): {temp_pw}")
                print("User must change password at first login.")
            else:
                print("Failed to create account.")
            pause()
        elif c == "3":

            clear_screen()
            print_header("Reset User Password", "Reason is required. This action will be logged.")
            username = ask_non_empty("Target username: ")
            u = store.find_user(username)
            if not u:
                print("User not found.")
                pause()
                continue

            reason = ask("Reason (required): ")
            if not reason.strip():
                print("❌ Reason is required. Reset canceled.")
                pause()
                continue

            new_pw = ask_password("New password: ")
            confirm = ask_password("Confirm new password: ")
            if not (new_pw or "").strip():
                print("❌ New password must not be empty.")
                pause()
                continue
            if new_pw != confirm:
                print("❌ New password does not match.")
                pause()
                continue

            ok_policy, msg = validate_password_policy(new_pw)
            if not ok_policy:
                print(f"❌ {msg}")
                pause()
                continue

            ok = store.update_password(username, new_pw)
            if ok:
                audit_log("RESET_PASSWORD", performed_by=me.username, target=username, reason=reason)
            print("✅ Password reset successfully! (Audit log saved)" if ok else "Reset failed.")
            pause()
        
        elif c == "4":
            clear_screen()
            print_header("Delete User", "Reason is required. This action will be logged. You cannot delete your own admin account.")
            username = ask_non_empty("Target username: ")

            # Safety checks
            if username.lower() == (me.username or "").lower():
                print("❌ You cannot delete the account you are currently using.")
                pause()
                continue

            u = store.find_user(username)
            if not u:
                print("User not found.")
                pause()
                continue

            # Prevent deleting the last Admin account
            if u.role == "Admin":
                admins = [x for x in store.list_users() if x.role == "Admin"]
                if len(admins) <= 1:
                    print("❌ Cannot delete the last Admin account.")
                    pause()
                    continue

            # ✅ Reason required (like reset password)
            reason = ask("Reason (required): ")
            if not reason.strip():
                print("❌ Reason is required. Delete canceled.")
                pause()
                continue

            if not ask_yes_no(f"Are you sure you want to delete '{u.username}' ({u.role})?", default=False):
                print("Canceled.")
                pause()
                continue

            ok = store.delete_user(u.username)
            if ok:
                audit_log("DELETE_ACCOUNT", performed_by=me.username, target=u.username, reason=reason)
                print("✅ Account deleted successfully! (Audit log saved)")
            else:
                print("Delete failed.")
            pause()

        


# -----------------------------
# Teacher
# -----------------------------
def teacher_menu(store: DataStore, me: User):
    while True:
        clear_screen()
        me = store.find_user(me.username) or me  # refresh profile
        subtitle = f"Logged in as: {me.full_name or me.username} | {me.username} | Role: Teacher"
        print_header("Teacher Panel", subtitle)
        print("1) Update profile")
        print("2) Change password")
        print("3) Template builder (create)")
        print("4) List my templates")
        print("5) Publish exam from template")
        print("6) List my exams")
        print("0) Logout")
        c = ask("Choose: ")

        if c == "0":
            return
        if c == "1":
            teacher_update_profile(store, me)
        elif c == "2":
            change_password_console(store, me)
            me = store.find_user(me.username) or me
        elif c == "3":
            teacher_create_template(store, me)
        elif c == "4":
            teacher_list_templates(store, me)
        elif c == "5":
            teacher_publish_exam(store, me)
        elif c == "6":
            teacher_list_exams(store, me)
        else:
            print("Invalid choice.")
            pause()


def teacher_update_profile(store: DataStore, me: User):
    clear_screen()
    print_header("Update Teacher Profile", "Leave blank to keep current value.")
    cur = store.find_user(me.username) or me

    full_name = ask(f"Full name [{cur.full_name}]: ")
    dob = ask(f"DOB [{cur.dob}] (YYYY-MM-DD): ")

    if not full_name:
        full_name = cur.full_name

    if not dob:
        dob = cur.dob
    else:
        if not is_valid_date_yyyy_mm_dd(dob):
            print("Invalid DOB. Update canceled.")
            pause()
            return

    # Teacher has no student_id
    ok = store.update_profile(me.username, full_name, dob, student_id="")
    print("Updated." if ok else "Update failed.")
    pause()


def teacher_create_template(store: DataStore, me: User):
    clear_screen()
    print_header("Create Template", "Each question has 4 options. Multi-correct is allowed.")
    title = ask_non_empty("Template title: ")

    qs: List[Question] = []
    while True:
        print("\n--- Add question ---")
        q_text = ask_non_empty("Question text: ")
        options = []
        for i in range(4):
            options.append(ask_non_empty(f"Option {i+1}: "))
        print("Mark correct options (example: 1 3). At least one is needed.")
        corr = ask_multi_choice(4)
        if not corr:
            print("Need at least 1 correct option. Question not added.")
            continue
        qs.append(Question(text=q_text, options=options, correct_indices=sorted(corr)))
        print("Question added.")
        if not ask_yes_no("Add another question?", default=True):
            break

    t = Template(template_id=store.new_template_id(), title=title, created_by=me.username, questions=qs)
    store.add_template(t)
    print(f"Template saved. ID: {t.template_id} | Questions: {len(qs)}")
    pause()


def teacher_list_templates(store: DataStore, me: User):
    while True:
        clear_screen()
        my = store.list_templates_by_teacher(me.username)
        print_header("My Templates")
        if not my:
            print("(No templates yet)")
            pause()
            return
        for i, t in enumerate(my, start=1):
            print(f"{i}) {t.template_id} | {t.title} | questions:{len(t.questions)}")
        print("\nOptions:")
        print("p) Preview  d) Delete  w) Export Word  0) Back")
        act = ask("Choose action: ").lower()
        if act == "0":
            return
        idx = ask_int("Template number: ", 1, len(my)) - 1
        t = my[idx]

        if act == "p":
            clear_screen()
            print_header(f"Template Preview: {t.title}", f"id:{t.template_id} | by:{t.created_by}")
            for qi, q in enumerate(t.questions, start=1):
                print("-" * 70)
                print(f"Q{qi}: {q.text}")
                for oi, opt in enumerate(q.options, start=1):
                    mark = " (correct)" if (oi - 1) in set(q.correct_indices) else ""
                    print(f"  {oi}. {opt}{mark}")
                print("Correct:", ", ".join(str(x + 1) for x in sorted(q.correct_indices)))
            pause()
        elif act == "d":
            if store.has_exam_from_template(t.template_id):
                print("This template has published exams. Deleting template will not delete exams.")
            if ask_yes_no(f"Delete template '{t.title}'?", default=False):
                ok = store.delete_template(t.template_id)
                print("Deleted." if ok else "Delete failed.")
                pause()
        elif act == "w":
            include_ans = ask_yes_no("Include correct answers in Word?", default=True)
            filepath = ask_non_empty("Output .docx path (example: template.docx): ")
            export_template_to_word(t, filepath, include_answers=include_ans)
            pause()
        else:
            print("Unknown action.")
            pause()


def teacher_publish_exam(store: DataStore, me: User):
    clear_screen()
    my = store.list_templates_by_teacher(me.username)
    print_header("Publish Exam", "Select a template to publish an exam (code will be generated).")
    if not my:
        print("No templates. Please create template first.")
        pause()
        return
    for i, t in enumerate(my, start=1):
        print(f"{i}) {t.template_id} | {t.title} | questions:{len(t.questions)}")
    idx = ask_int("Choose template number: ", 1, len(my)) - 1
    tpl = my[idx]

    start_ts = ask_dt("Start")
    end_ts = ask_dt("End")
    if end_ts <= start_ts:
        print("End must be after Start.")
        pause()
        return

    dur_min = ask_int("Duration (minutes, 1..240): ", 1, 240)
    attempt_limit = ask_int("Attempt limit (0..50, 0=unlimited): ", 0, 50)
    allow_review = ask_yes_no("Allow review after submit?", default=False)

    use_pass = ask_yes_no("Use password for this exam?", default=False)
    password = ""
    if use_pass:
        password = ask_non_empty("Exam password: ")

    code = store.new_unique_code(8)
    e = Exam(
        exam_id=store.new_exam_id(),
        template_id=tpl.template_id,
        title=tpl.title,
        created_by=me.username,
        access_code=code,
        password=password,
        duration_seconds=dur_min * 60,
        allow_review=allow_review,
        attempt_limit=attempt_limit,
        start_ts=int(start_ts),
        end_ts=int(end_ts),
        questions=list(tpl.questions),
    )
    store.add_exam(e)
    print(f"Published exam OK. CODE: {e.access_code} | Exam ID: {e.exam_id}")
    pause()


def teacher_list_exams(store: DataStore, me: User):
    while True:
        clear_screen()
        now = int(time.time())
        my = store.list_exams_by_teacher(me.username)
        print_header("My Exams")
        if not my:
            print("(No exams yet)")
            pause()
            return

        for i, e in enumerate(my, start=1):
            status = "OPEN" if (e.start_ts <= now <= e.end_ts) else ("WAIT" if now < e.start_ts else "CLOSED")
            mins = max(1, e.duration_seconds // 60)
            print(f"{i}) {e.exam_id} | {e.title} | code:{e.access_code} | {status} | {fmt_dt(e.start_ts)} -> {fmt_dt(e.end_ts)} | {mins} min | limit:{e.attempt_limit}")

        print("\nOptions:")
        print("p) Preview  a) Attempts/results  w) Export Word  x) Export Excel  d) Delete exam  r) Delete attempts  0) Back")
        act = ask("Choose action: ").lower()
        if act == "0":
            return
        idx = ask_int("Exam number: ", 1, len(my)) - 1
        e = my[idx]

        if act == "p":
            teacher_preview_exam(e)
        elif act == "a":
            teacher_exam_attempts_menu(store, e)
        elif act == "w":
            include_ans = ask_yes_no("Include correct answers in Word?", default=True)
            filepath = ask_non_empty("Output .docx path (example: exam.docx): ")
            export_exam_to_word(e, filepath, include_answers=include_ans)
            pause()
        elif act == "x":
            filepath = ask_non_empty("Output .xlsx path (example: results.xlsx): ")
            attempts = store.list_attempts_for_exam(e.exam_id)
            export_exam_results_to_excel(e, attempts, filepath)
            pause()
        elif act == "d":
            if ask_yes_no(f"Delete exam '{e.title}' (attempts will remain)?", default=False):
                ok = store.delete_exam(e.exam_id)
                print("Deleted." if ok else "Delete failed.")
                pause()
        elif act == "r":
            if ask_yes_no("Delete ALL attempts for this exam?", default=False):
                deleted = store.delete_attempts_for_exam(e.exam_id)
                print(f"Deleted attempts: {deleted}")
                pause()
        else:
            print("Unknown action.")
            pause()


def teacher_preview_exam(e: Exam):
    clear_screen()
    print_header(f"Exam Preview: {e.title}", f"id:{e.exam_id} | code:{e.access_code}")
    print(f"Window: {fmt_dt(e.start_ts)} -> {fmt_dt(e.end_ts)}")
    print(f"Duration: {max(1, e.duration_seconds // 60)} minutes")
    print(f"Password: {'set' if e.password else 'none'}")
    print(f"Allow review: {'yes' if e.allow_review else 'no'}")
    print(f"Attempt limit: {e.attempt_limit} (0=unlimited)")
    print(f"From template: {e.template_id}")
    for i, q in enumerate(e.questions, start=1):
        print("-" * 70)
        print(f"Q{i}: {q.text}")
        for oi, opt in enumerate(q.options, start=1):
            mark = " (correct)" if (oi - 1) in set(q.correct_indices) else ""
            print(f"  {oi}. {opt}{mark}")
    pause()


def teacher_exam_attempts_menu(store: DataStore, e: Exam):
    while True:
        clear_screen()
        attempts = store.list_attempts_for_exam(e.exam_id)
        print_header("Attempts / Results", f"{e.title} | code:{e.access_code}")
        if not attempts:
            print("(No attempts yet)")
            pause()
            return

        total_q = max(1, len(e.questions))
        for i, a in enumerate(attempts, start=1):
            score10 = (a.score / max(1, a.total)) * 10.0
            took = f"{a.time_taken_seconds // 60:02d}:{a.time_taken_seconds % 60:02d}"
            print(f"{i}) {a.full_name or '(no name)'} | {a.username} | {score10:.2f}/10 | took {took} | {fmt_dt_full(a.submitted_at)}")

        print("\nOptions:")
        print("v) View details  w) Export attempt Word  0) Back")
        act = ask("Choose: ").lower()
        if act == "0":
            return
        idx = ask_int("Attempt number: ", 1, len(attempts)) - 1
        a = attempts[idx]
        if act == "v":
            teacher_view_attempt_details(e, a)
        elif act == "w":
            filepath = ask_non_empty("Output .docx path (example: attempt.docx): ")
            export_attempt_to_word(e, a, filepath)
            pause()
        else:
            print("Unknown action.")
            pause()


def teacher_view_attempt_details(e: Exam, a: Attempt):
    clear_screen()
    total_q = max(1, len(e.questions))
    score10 = (a.score / total_q) * 10.0
    print_header("Attempt Details", f"{e.title} | {a.full_name} | {a.username} | {score10:.2f}/10")
    print(f"Started:   {fmt_dt_full(a.started_at)}")
    print(f"Submitted: {fmt_dt_full(a.submitted_at)}")
    print(f"Time taken: {a.time_taken_seconds} sec")
    points_per_q = 10.0 / total_q
    answers = a.answers or []

    for i, q in enumerate(e.questions):
        user_sel = set(answers[i]) if i < len(answers) else set()
        correct = set(q.correct_indices)
        earned, missing, extra = score_question_partial(user_sel, correct, points=points_per_q)
        print("-" * 70)
        print(f"Q{i+1}: {q.text}")
        print(f"Earned: {earned:.2f}/{points_per_q:.2f}")
        for oi, opt in enumerate(q.options, start=1):
            mu = "[x]" if (oi - 1) in user_sel else "[ ]"
            mc = " (correct)" if (oi - 1) in correct else ""
            print(f"  {mu} {oi}. {opt}{mc}")
        print("Missing correct:", (", ".join(str(x + 1) for x in sorted(missing)) if missing else "none"))
        print("Extra wrong:", (", ".join(str(x + 1) for x in sorted(extra)) if extra else "none"))
    pause()


def change_password_console(store: DataStore, me: User):
    clear_screen()
    print_header("Change Password")

    current = ask_password("Current password: ")
    cur = store.find_user(me.username)
    if (not cur) or ((cur.password or "") != (current or "")):
        print("❌ Current password is incorrect.")
        pause()
        return

    new_pw = ask_password("New password: ")
    confirm = ask_password("Confirm new password: ")

    if not (new_pw or "").strip():
        print("❌ New password must not be empty.")
        pause()
        return
    if new_pw != confirm:
        print("❌ New password does not match.")
        pause()
        return

    ok_policy, msg = validate_password_policy(new_pw)
    if not ok_policy:
        print(f"❌ {msg}")
        pause()
        return

    ok = store.update_password(me.username, new_pw)
    if not ok:
        print("❌ Failed to change password.")
        pause()
        return

    # Clear the first-login requirement
    store.set_must_change_password(me.username, False)
    me.must_change_password = False
    print("✅ Password changed successfully!")


# -----------------------------
# Student
# -----------------------------
def student_menu(store: DataStore, me: User):
    while True:
        clear_screen()
        print_header("Student Panel", "Join exams by CODE. Update your profile and view history.")
        print("1) Update profile")
        print("2) Change password")
        print("3) Enter exam by CODE")
        print("4) My attempts")
        print("0) Logout")
        c = ask("Choose: ")
        if c == "0":
            return
        if c == "1":
            student_update_profile(store, me)
            me = store.find_user(me.username) or me
        elif c == "2":
            change_password_console(store, me)
            me = store.find_user(me.username) or me
        elif c == "3":
            student_enter_exam(store, me)
        elif c == "4":
            student_my_attempts(store, me)
        else:
            print("Invalid choice.")
            pause()


def student_update_profile(store: DataStore, me: User):
    clear_screen()
    print_header("Update Profile", "Leave blank to keep current value.")
    cur = store.find_user(me.username) or me
    full_name = ask(f"Full name [{cur.full_name}]: ")
    dob = ask(f"DOB [{cur.dob}] (YYYY-MM-DD): ")
    sid = ask(f"Student ID [{cur.student_id}]: ")

    if not full_name:
        full_name = cur.full_name
    if not dob:
        dob = cur.dob
    else:
        if not is_valid_date_yyyy_mm_dd(dob):
            print("Invalid DOB. Update canceled.")
            pause()
            return
    if not sid:
        sid = cur.student_id

    ok = store.update_profile(me.username, full_name, dob, sid)
    print("Updated." if ok else "Update failed.")
    pause()


def _check_exam_password(exam: Exam) -> bool:
    if not exam.password:
        return True
    pw = ask("Exam password: ")
    if pw != exam.password:
        print("Password is not correct.")
        pause()
        return False
    return True


def student_enter_exam(store: DataStore, me: User):
    clear_screen()
    print_header("Enter Exam")
    code = ask_non_empty("CODE: ").upper()
    exam = store.get_exam_by_code(code)
    if not exam:
        print("Exam not found by this code.")
        pause()
        return

    now = int(time.time())
    if now < exam.start_ts:
        print(f"Exam not open yet. Start: {fmt_dt(exam.start_ts)}")
        pause()
        return
    if now > exam.end_ts:
        print(f"Exam closed. End: {fmt_dt(exam.end_ts)}")
        pause()
        return

    if exam.attempt_limit > 0:
        used = store.count_attempts_for_user_exam(me.username, exam.exam_id)
        if used >= exam.attempt_limit:
            print(f"Attempt limit reached. Limit: {exam.attempt_limit}, Used: {used}")
            pause()
            return

    if not _check_exam_password(exam):
        return

    student_take_exam(store, me, exam)


def student_take_exam(store: DataStore, me: User, exam: Exam):
    clear_screen()
    mins = max(1, exam.duration_seconds // 60)
    print_header("Take Exam", f"{exam.title} | code:{exam.access_code} | {mins} min")
    print("Rule: for multi-correct questions, partial credit may apply.")
    print("Note: Console demo checks time only at each question input.\n")

    answers: List[Set[int]] = [set() for _ in range(len(exam.questions))]
    started_at = time.time()
    end_time = started_at + exam.duration_seconds

    for i, q in enumerate(exam.questions):
        clear_screen()
        left = int(end_time - time.time())
        if left <= 0:
            print("Time is over. Auto submit now.")
            break

        mm, ss = max(0, left) // 60, max(0, left) % 60
        print_header("Take Exam", f"Time left: {mm:02d}:{ss:02d} | Q{i+1}/{len(exam.questions)}")
        print(q.text)
        for oi, opt in enumerate(q.options, start=1):
            print(f"  {oi}. {opt}")
        answers[i] = ask_multi_choice(4)

    submitted_at = time.time()
    time_taken = int(max(0, submitted_at - started_at))

    total_q = max(1, len(exam.questions))
    total_score = 0.0
    for ans, q in zip(answers, exam.questions):
        earned, _, _ = score_question_partial(ans, set(q.correct_indices), points=1.0)
        total_score += earned

    # refresh user profile data
    u = store.find_user(me.username) or me

    a = Attempt(
        attempt_id=store.new_attempt_id(),
        exam_id=exam.exam_id,
        code=exam.access_code,
        title=exam.title,
        username=u.username,
        full_name=u.full_name,
        student_id=u.student_id,
        score=total_score,
        total=total_q,
        started_at=started_at,
        submitted_at=submitted_at,
        time_taken_seconds=time_taken,
        answers=[sorted(list(s)) for s in answers],
    )
    store.add_attempt(a)

    score10 = (total_score / total_q) * 10.0
    clear_screen()
    print_header("Submitted")
    print(f"Score: {score10:.2f}/10 (raw {total_score:.4f}/{total_q})")
    print(f"Time taken: {time_taken} sec")
    print(f"Submitted at: {fmt_dt_full(submitted_at)}")
    pause()


def student_my_attempts(store: DataStore, me: User):
    while True:
        clear_screen()
        attempts = store.list_attempts_for_user(me.username)
        print_header("My Attempts")
        if not attempts:
            print("No attempts yet.")
            pause()
            return

        for i, a in enumerate(attempts, start=1):
            score10 = (a.score / max(1, a.total)) * 10.0
            print(f"{i}) {fmt_dt_full(a.submitted_at)} | {a.title} | code:{a.code} | {score10:.2f}/10")

        print("\nOptions:")
        print("r) Review attempt (if allowed)  0) Back")
        act = ask("Choose: ").lower()
        if act == "0":
            return
        if act != "r":
            print("Unknown action.")
            pause()
            continue
        idx = ask_int("Attempt number: ", 1, len(attempts)) - 1
        attempt = attempts[idx]
        exam = store.get_exam(attempt.exam_id)
        if not exam:
            print("Exam data not found.")
            pause()
            continue
        if not exam.allow_review:
            print("Teacher does not allow review for this exam.")
            pause()
            continue
        student_review_attempt(exam, attempt)


def student_review_attempt(exam: Exam, attempt: Attempt):
    clear_screen()
    total_q = max(1, len(exam.questions))
    score10 = (attempt.score / total_q) * 10.0
    need_more = max(0.0, 10.0 - score10)
    took = f"{attempt.time_taken_seconds // 60:02d}:{attempt.time_taken_seconds % 60:02d}"
    print_header("Review", f"{exam.title} | {attempt.username} | {score10:.2f}/10 | need {need_more:.2f} | took {took}")
    points_per_q = 10.0 / total_q
    answers = attempt.answers or []

    for i, q in enumerate(exam.questions):
        user_sel = set(answers[i]) if i < len(answers) else set()
        correct = set(q.correct_indices)
        earned, missing, extra = score_question_partial(user_sel, correct, points=points_per_q)
        print("-" * 70)
        print(f"Q{i+1}: {q.text}")
        print(f"Earned: {earned:.2f}/{points_per_q:.2f}")
        for oi, opt in enumerate(q.options, start=1):
            mu = "[x]" if (oi - 1) in user_sel else "[ ]"
            mc = " (correct)" if (oi - 1) in correct else ""
            print(f"  {mu} {oi}. {opt}{mc}")
        print("Missing correct:", (", ".join(str(x + 1) for x in sorted(missing)) if missing else "none"))
        print("Extra wrong:", (", ".join(str(x + 1) for x in sorted(extra)) if extra else "none"))
    pause()


# -----------------------------
# Main loop
# -----------------------------
def main():
    store = DataStore(DATA_FILE)

    while True:
        u = login_flow(store)
        if u is None:
            continue
        if u.username == "__EXIT__":
            break

        if u.role == "Admin":
            admin_menu(store, u)
        elif u.role == "Teacher":
            teacher_menu(store, u)
        else:
            student_menu(store, u)


if __name__ == "__main__":
    main()
