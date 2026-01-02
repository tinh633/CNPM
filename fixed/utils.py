# utils.py
import time
import re
import secrets
import string
from tkinter import messagebox
from typing import List, Optional, Set, Tuple

from models import Template, Exam, Attempt, score_question_partial

try:
    from docx import Document
    from openpyxl import Workbook
except ImportError:
    pass


# -----------------------------
# Password + Username policies (GUI requirements)
# -----------------------------
_PASSWORD_SPECIALS = "!@#$%^&*()-_=+[]{};:,.?/\\|~`"

def is_valid_cccd_username(username: str) -> bool:
    """Username must be exactly 12 digits (CCCD)."""
    u = (username or "").strip()
    return bool(re.fullmatch(r"\d{12}", u))

def validate_temp_password(pw: str) -> Tuple[bool, str]:
    """Temp password: exactly 8+ chars is OK for temp; we will generate 8 exactly for admin-created accounts."""
    if pw is None:
        return False, "Password is required."
    if len(pw) < 8:
        return False, "Password must be at least 8 characters."
    if not re.search(r"[A-Z]", pw):
        return False, "Password must contain at least one uppercase letter."
    if not re.search(r"[a-z]", pw):
        return False, "Password must contain at least one lowercase letter."
    if not re.search(r"\d", pw):
        return False, "Password must contain at least one digit."
    if not re.search(rf"[{re.escape(_PASSWORD_SPECIALS)}]", pw):
        return False, f"Password must contain at least one special character ({_PASSWORD_SPECIALS})."
    return True, "OK"

def validate_strong_password(pw: str) -> Tuple[bool, str]:
    """Strong password for user change: >8 chars and contains upper/lower/digit/special."""
    if pw is None:
        return False, "Password is required."
    if len(pw) <= 8:
        return False, "Password must be more than 8 characters."
    return validate_temp_password(pw)

def generate_temp_password(length: int = 8) -> str:
    """Generate temp password with required categories. Default length = 8."""
    length = int(length)
    if length < 8:
        length = 8
    # ensure all categories
    upp = secrets.choice(string.ascii_uppercase)
    low = secrets.choice(string.ascii_lowercase)
    dig = secrets.choice(string.digits)
    spe = secrets.choice(_PASSWORD_SPECIALS)
    rest = [secrets.choice(string.ascii_letters + string.digits + _PASSWORD_SPECIALS) for _ in range(length - 4)]
    pw_list = [upp, low, dig, spe] + rest
    secrets.SystemRandom().shuffle(pw_list)
    pw = "".join(pw_list)
    ok, _ = validate_temp_password(pw)
    return pw if ok else generate_temp_password(length)

# -----------------------------
# Time helpers
# -----------------------------
def parse_dt(s: str) -> Optional[int]:
    """Parse 'YYYY-MM-DD HH:MM' to unix timestamp (local)."""
    s = (s or "").strip()
    try:
        return int(time.mktime(time.strptime(s, "%Y-%m-%d %H:%M")))
    except Exception:
        return None

def fmt_dt(ts: int) -> str:
    try:
        return time.strftime("%Y-%m-%d %H:%M", time.localtime(int(ts)))
    except Exception:
        return "N/A"

def fmt_dt_full(ts: float) -> str:
    try:
        return time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(float(ts)))
    except Exception:
        return "N/A"

# -----------------------------
# Export helpers
# -----------------------------
def export_template_to_word(t: Template, filepath: str, include_answers: bool = True):
    try:
        doc = Document()
        doc.add_heading(f"Template: {t.title}", level=1)
        doc.add_paragraph(f"Template ID: {t.template_id}")
        doc.add_paragraph(f"Created by: {t.created_by}")
        doc.add_paragraph(f"Questions: {len(t.questions)}")

        for i, q in enumerate(t.questions, start=1):
            doc.add_heading(f"Q{i}: {q.text}", level=2)
            for oi, opt in enumerate(q.options, start=1):
                doc.add_paragraph(f"{oi}. {opt}", style="List Bullet")
            if include_answers:
                correct = ", ".join(str(x + 1) for x in sorted(q.correct_indices))
                doc.add_paragraph(f"Correct: {correct}")
        doc.save(filepath)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export docx: {e}")

def export_exam_results_to_excel(e: Exam, attempts: List[Attempt], filepath: str):
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Results"

        ws.append([
            "Exam Title", "Exam ID", "Code",
            "Username", "Full Name", "Student ID",
            "Score (/10)", "Raw Score", "Total Questions",
            "Time Taken (sec)", "Submitted At"
        ])

        for a in attempts:
            score10 = (a.score / max(1, a.total)) * 10.0
            ws.append([
                e.title, e.exam_id, e.access_code,
                a.username, a.full_name, a.student_id,
                round(score10, 2), round(a.score, 4), a.total,
                a.time_taken_seconds, fmt_dt_full(a.submitted_at)
            ])
        wb.save(filepath)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export xlsx: {e}")

# -----------------------------
# Audit / Security helpers
# -----------------------------
def audit_log(action: str, performed_by: str, target: str = "", reason: str = ""):
    """
    Security audit log (non-functional requirement).
    """
    ts = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
    line = f"[{ts}] action={action} by={performed_by}"
    if target:
        line += f" target={target}"
    if reason:
        line += f" reason={reason.replace(chr(10), ' ')}"
    line += "\n"

    try:
        with open("audit.log", "a", encoding="utf-8") as f:
            f.write(line)
    except Exception:
        pass